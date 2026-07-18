#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path
from typing import Any

# Compatibilidade temporária entre pacote oficial e execução direta.
if __package__:
    from .prisma_model import PrismaReport, StudyRecord
else:
    from prisma_model import PrismaReport, StudyRecord

try:
    from docx import Document
    from docx.shared import Cm, Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:  # pragma: no cover
    Document = None  # type: ignore


def setup_docx_styles(doc: Any) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for name, size in (("Title", 16), ("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)):
        try:
            st = doc.styles[name]
            st.font.name = "Times New Roman"
            st.font.size = Pt(size)
            st.font.bold = True
        except Exception:
            pass


def add_justified_paragraph(doc: Any, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_table(doc: Any, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        for run in hdr[idx].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value or "")
    doc.add_paragraph("")


def study_rows(studies: list[StudyRecord]) -> list[list[str]]:
    rows: list[list[str]] = []
    for i, study in enumerate(studies, start=1):
        rows.append([
            str(i), study.titulo, "; ".join(study.autores), study.ano,
            study.doi, study.bib_key, study.justificativa or study.motivo,
        ])
    return rows


def render_prisma_docx(report: PrismaReport, output_path: Path, reference_docx: Path | None = None, flow_svg_path: Path | None = None) -> Path:
    if Document is None:
        raise RuntimeError("python-docx não está instalado.")
    doc = Document(str(reference_docx)) if reference_docx and reference_docx.exists() else Document()
    setup_docx_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(3)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    sec.bottom_margin = Cm(2)

    meta = report.metadata
    for line in [meta.instituicao, meta.curso, meta.disciplina, f"Professor: {meta.professor}" if meta.professor else "", meta.responsavel]:
        if line:
            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(meta.titulo)
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("")
    p = doc.add_paragraph(f"{meta.cidade}, {meta.data_execucao}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("1 Introdução do relatório", level=1)
    add_justified_paragraph(doc, "Este relatório apresenta a trilha metodológica de identificação, triagem, elegibilidade e inclusão dos textos utilizados na pesquisa. O objetivo é registrar de forma auditável as bases ou fontes consultadas, os critérios de inclusão e exclusão, os totais do fluxo e os estudos efetivamente incorporados ao corpus.")

    doc.add_heading("2 Tema, recorte e objetivo da busca", level=1)
    for label, value in [("Tema", meta.tema), ("Recorte", meta.recorte), ("Objetivo", meta.objetivo), ("Pergunta de pesquisa", meta.pergunta_pesquisa)]:
        doc.add_paragraph(f"{label}: {value or 'Não informado'}")

    doc.add_heading("3 Estratégia de busca", level=1)
    doc.add_paragraph("Bases/fontes: " + ("; ".join(report.search_strategy.bases) if report.search_strategy.bases else "Não informado"))
    doc.add_paragraph("Idiomas: " + ("; ".join(report.search_strategy.idiomas) if report.search_strategy.idiomas else "Não informado"))
    doc.add_paragraph("Período: " + (report.search_strategy.periodo or "Não informado"))
    if report.search_strategy.queries:
        add_table(doc, ["Base", "Consulta", "Resultados brutos", "Observações"], [[q.base, q.query, str(q.resultados_brutos), q.observacoes] for q in report.search_strategy.queries])

    doc.add_heading("4 Critérios de inclusão e exclusão", level=1)
    doc.add_heading("4.1 Critérios de inclusão", level=2)
    for item in report.criteria.inclusao:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("4.2 Critérios de exclusão", level=2)
    for item in report.criteria.exclusao:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("5 Resultados da busca e da triagem", level=1)
    add_table(doc, ["Etapa", "Total"], [
        ["Registros identificados", str(report.flow.identificados)],
        ["Duplicados removidos", str(report.flow.duplicados_removidos)],
        ["Registros após deduplicação", str(report.flow.apos_deduplicacao)],
        ["Registros triados por título/resumo", str(report.flow.triados_titulo_resumo)],
        ["Excluídos na triagem", str(report.flow.excluidos_titulo_resumo)],
        ["Textos completos avaliados", str(report.flow.avaliados_texto_completo)],
        ["Excluídos após texto completo", str(report.flow.excluidos_texto_completo)],
        ["Estudos incluídos", str(report.flow.incluidos)],
    ])

    doc.add_heading("6 Fluxograma PRISMA", level=1)
    if flow_svg_path and flow_svg_path.exists():
        try:
            doc.add_picture(str(flow_svg_path), width=Inches(5.8))
        except Exception:
            # python-docx não insere SVG em alguns ambientes; a tabela acima permanece como representação textual do fluxo.
            doc.add_paragraph("Fluxograma PRISMA registrado em arquivo externo: " + str(flow_svg_path))
    else:
        doc.add_paragraph("Fluxograma PRISMA representado pelos totais da seção anterior.")

    doc.add_heading("7 Estudos incluídos", level=1)
    if report.included_studies:
        add_table(doc, ["#", "Título", "Autores", "Ano", "DOI", "Chave", "Justificativa"], study_rows(report.included_studies))
    else:
        doc.add_paragraph("Nenhum estudo incluído registrado.")

    doc.add_heading("8 Estudos excluídos", level=1)
    if report.excluded_studies:
        add_table(doc, ["#", "Título", "Autores", "Ano", "DOI", "Chave", "Motivo"], study_rows(report.excluded_studies))
    else:
        doc.add_paragraph("Nenhum estudo excluído registrado.")

    doc.add_heading("9 Diagnóstico metodológico", level=1)
    doc.add_paragraph("Fontes de artefatos: " + ("; ".join(report.diagnostics.fontes_artefatos) if report.diagnostics.fontes_artefatos else "Não informado"))
    doc.add_paragraph("Bases com erro: " + ("; ".join(report.diagnostics.bases_com_erro) if report.diagnostics.bases_com_erro else "Nenhuma registrada"))
    doc.add_paragraph("Avisos: " + ("; ".join(report.diagnostics.avisos) if report.diagnostics.avisos else "Nenhum"))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
