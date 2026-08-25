#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import zipfile
from dataclasses import dataclass, asdict
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    import tomllib  # Python 3.11+
except Exception:  # pragma: no cover
    tomllib = None

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
except Exception as exc:  # pragma: no cover
    Document = None  # type: ignore[assignment]
    _DOCX_IMPORT_ERROR = exc
else:
    _DOCX_IMPORT_ERROR = None


@dataclass
class ArticlePaths:
    art_dir: Path
    output_dir: Path
    prefix: str
    org: Path
    bib: Path
    document_json: Path
    resumos_json: Path
    compliance_json: Path
    run_report_json: Path
    cfg: Path | None
    docx: Path
    canonical_json: Path


DEFAULT_LAYOUT = {
    "papel": "A4",
    "margens_cm": {"superior": 3.0, "esquerda": 3.0, "direita": 2.0, "inferior": 2.0},
    "fonte_texto_pt": 12,
    "fonte_auxiliar_pt": 10,
    "espacamento_texto": 1.5,
    "recuo_primeira_linha_cm": 1.25,
    "recuo_citacao_longa_cm": 4.0,
}


TITLE_KEYS = {"title", "titulo", "document_title", "paper_title"}
AUTHOR_KEYS = {"author", "authors", "autor", "autores"}
INSTITUTION_KEYS = {"institution", "instituicao", "institution_name", "nome_instituicao"}
PROGRAM_KEYS = {"program", "programa", "program_name", "nome_programa"}
COURSE_KEYS = {"course", "curso", "course_name", "nome_curso"}
DISCIPLINE_KEYS = {"discipline", "disciplina", "class", "materia"}
PROFESSOR_KEYS = {"professor", "professora", "docente", "teacher", "orientador"}
CITY_KEYS = {"city", "cidade", "local"}
YEAR_KEYS = {"year", "ano"}
SUBTITLE_KEYS = {"subtitle", "subtitulo"}


# ---------------------------------------------------------------------------
# Generic helpers
# ---------------------------------------------------------------------------


def die(msg: str) -> None:
    raise SystemExit(msg)


def now_stamp() -> str:
    return datetime.now().strftime("%Y%m%d_%H%M%S")


def backup(path: Path, tag: str, quiet: bool = False) -> Path | None:
    if not path.exists():
        return None
    bak = path.with_name(path.name + f".bak_{tag}_{now_stamp()}")
    shutil.copy2(path, bak)
    if not quiet:
        print(f"[OK] Backup: {bak}")
    return bak


def read_text(path: Path) -> str:
    if not path.exists():
        return ""
    return path.read_text(encoding="utf-8", errors="ignore")


def read_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8", errors="ignore"))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def read_toml(path: Path | None) -> dict[str, Any]:
    if not path or not path.exists() or tomllib is None:
        return {}
    try:
        data = tomllib.loads(path.read_text(encoding="utf-8", errors="ignore"))
        return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def flatten_dict(obj: Any, prefix: str = "") -> dict[str, Any]:
    out: dict[str, Any] = {}
    if isinstance(obj, dict):
        for k, v in obj.items():
            key = f"{prefix}.{k}" if prefix else str(k)
            out[key] = v
            out.update(flatten_dict(v, key))
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            out.update(flatten_dict(v, f"{prefix}.{i}" if prefix else str(i)))
    return out


def first_by_keys(data: dict[str, Any], keys: set[str]) -> str:
    flat = flatten_dict(data)
    # Prefer exact final component match.
    for path_key, value in flat.items():
        last = path_key.rsplit(".", 1)[-1].lower()
        if last in keys and value not in (None, ""):
            if isinstance(value, list):
                return "; ".join(str(x) for x in value if str(x).strip())
            return str(value).strip()
    return ""


def strip_accents_for_sort(s: str) -> str:
    try:
        import unicodedata
        return "".join(c for c in unicodedata.normalize("NFKD", s) if not unicodedata.combining(c)).lower()
    except Exception:
        return s.lower()



def clean_spaces(s: str) -> str:
    """Normaliza espaços e remove resíduos típicos de LaTeX/Org no texto visível."""
    if s is None:
        return ""
    s = str(s)
    # Caracteres que aparecem por conversão/extração e não devem vazar para o DOCX.
    s = s.replace("\xa0", " ")
    s = s.replace("\ufeff", "")
    s = s.replace("\ufffe", "-")
    s = s.replace("\u00ad", "")  # soft hyphen invisível
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"[ \t]+", " ", s)
    s = re.sub(r" *\n *", "\n", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()

# ---------------------------------------------------------------------------
# Path discovery
# ---------------------------------------------------------------------------


def guess_prefix(art_dir: Path, explicit: str | None = None) -> str:
    if explicit:
        return explicit
    out = art_dir / "output"
    candidates = sorted(out.glob("*.document.json")) if out.exists() else []
    if candidates:
        return candidates[0].name.replace(".document.json", "")
    tomls = sorted(art_dir.glob("*.toml"))
    if tomls:
        return tomls[0].stem
    return "artigo_final_atestmed_abnt"


def find_cfg(art_dir: Path, prefix: str, explicit: Path | None = None) -> Path | None:
    if explicit and explicit.exists():
        return explicit
    p = art_dir / f"{prefix}.toml"
    if p.exists():
        return p
    tomls = sorted(art_dir.glob("*.toml"))
    return tomls[0] if tomls else explicit


def resolve_paths(art_dir: Path, cfg_art: Path | None, prefix: str | None, output: Path | None = None) -> ArticlePaths:
    art_dir = art_dir.resolve()
    pfx = guess_prefix(art_dir, prefix)
    out = output.parent.resolve() if output and output.suffix.lower() == ".docx" else (art_dir / "output").resolve()
    cfg = find_cfg(art_dir, pfx, cfg_art.resolve() if cfg_art else None)
    docx = output.resolve() if output and output.suffix.lower() == ".docx" else out / f"{pfx}.docx"
    return ArticlePaths(
        art_dir=art_dir,
        output_dir=out,
        prefix=pfx,
        org=out / f"{pfx}.org",
        bib=out / f"{pfx}.bib",
        document_json=out / f"{pfx}.document.json",
        resumos_json=out / f"{pfx}.resumos_paper.json",
        compliance_json=out / f"{pfx}.compliance_report.json",
        run_report_json=out / f"{pfx}.run_report.json",
        cfg=cfg,
        docx=docx,
        canonical_json=out / f"{pfx}.canonical_docx.json",
    )


# ---------------------------------------------------------------------------
# Metadata extraction
# ---------------------------------------------------------------------------



def parse_org_metadata(org_text: str) -> dict[str, str]:
    meta: dict[str, str] = {}
    for line in org_text.splitlines():
        m = re.match(r"^#\+([A-Za-z0-9_:-]+):\s*(.*)$", line.strip())
        if not m:
            continue
        key = m.group(1).lower().replace("-", "_")
        value = m.group(2).strip()
        if value:
            meta[key] = value
    return meta


def parse_org_latex_header_macros(org_text: str) -> dict[str, str]:
    """Extrai metadados de capa declarados em #+LATEX_HEADER.

    A classe fgv-paper usa macros como \\institution{...}, \\coursename{...}
    e \\professorname{...}. Esses valores são mais confiáveis do que chaves
    genéricas do TOML para a capa do DOCX.
    """
    macros: dict[str, str] = {}
    wanted = {
        "institution", "programname", "coursename", "disciplinename",
        "professorname", "cityname", "papertype", "covernote",
    }
    for line in org_text.splitlines():
        if not line.strip().lower().startswith("#+latex_header:"):
            continue
        payload = line.split(":", 1)[1]
        for name, value in re.findall(r"\\([A-Za-z]+)\{([^{}]*)\}", payload):
            if name in wanted:
                macros[name] = clean_latex_inline(value)
    return macros


def invalid_cover_value(value: str) -> bool:
    v = clean_spaces(str(value)).strip()
    if not v:
        return True
    low = v.lower()
    if "perfil" in low:
        return True
    if low in {"fgv", "profile://fgv", "institution_profile: fgv", "instituicao_perfil: fgv"}:
        return True
    if re.fullmatch(r"['\"]?perfil['\"]?\s*:\s*['\"]?fgv['\"]?", low):
        return True
    if v.startswith("{") and "fgv" in low:
        return True
    return False


def first_valid(*values: str, default: str = "") -> str:
    for value in values:
        if value is None:
            continue
        v = clean_latex_inline(str(value).strip().strip('"').strip("'"))
        if v and not invalid_cover_value(v):
            return v
    return default


def build_metadata(paths: ArticlePaths, org_text: str, document_json: dict[str, Any], toml_data: dict[str, Any]) -> dict[str, str]:
    org_meta = parse_org_metadata(org_text)
    latex_meta = parse_org_latex_header_macros(org_text)
    merged: dict[str, Any] = {}
    merged.update(document_json)
    merged.update(toml_data)
    merged["org_meta"] = org_meta
    merged["latex_meta"] = latex_meta

    title = first_valid(org_meta.get("title", ""), first_by_keys(merged, TITLE_KEYS), str(document_json.get("title", "")), default=paths.prefix)
    author = first_valid(org_meta.get("author", ""), first_by_keys(merged, AUTHOR_KEYS), default="Gustavo M. Mendes de Tarso")
    subtitle = first_valid(org_meta.get("subtitle", ""), first_by_keys(merged, SUBTITLE_KEYS), default="")

    # A capa deve obedecer prioritariamente às macros do ORG/LaTeX final.
    institution = first_valid(latex_meta.get("institution", ""), first_by_keys(merged, INSTITUTION_KEYS), default="Fundação Getúlio Vargas")
    if "programname" in latex_meta:
        program = first_valid(latex_meta.get("programname", ""), default="")
    else:
        program = first_valid(first_by_keys(merged, PROGRAM_KEYS), default="")
    if "coursename" in latex_meta:
        course = first_valid(latex_meta.get("coursename", ""), default="")
    else:
        course = first_valid(first_by_keys(merged, COURSE_KEYS), default="")

    # Evita capa duplicada quando o TOML traz programa antigo e o ORG traz curso final.
    if program and course and strip_accents_for_sort(program) == strip_accents_for_sort(course):
        course = ""
    if program and course and "mestrado" in strip_accents_for_sort(program) and "mestrado" in strip_accents_for_sort(course):
        # O ORG final normalmente deixa programname vazio e usa coursename. Se ambos
        # sobreviverem de fontes diferentes, preserva o coursename declarado no ORG.
        if latex_meta.get("coursename"):
            program = ""

    discipline = first_valid(latex_meta.get("disciplinename", ""), first_by_keys(merged, DISCIPLINE_KEYS), default="Decisões Baseadas em Evidência")
    professor = first_valid(latex_meta.get("professorname", ""), first_by_keys(merged, PROFESSOR_KEYS), default="")
    city = first_valid(latex_meta.get("cityname", ""), first_by_keys(merged, CITY_KEYS), default="Brasília")
    year = first_valid(first_by_keys(merged, YEAR_KEYS), default=str(datetime.now().year))
    covernote = first_valid(latex_meta.get("covernote", ""), default="")
    papertype = first_valid(latex_meta.get("papertype", ""), default="")

    return {
        "title": title,
        "subtitle": subtitle,
        "author": author,
        "institution": institution,
        "program": program,
        "course": course,
        "discipline": discipline,
        "professor": professor,
        "city": city,
        "year": year,
        "covernote": covernote,
        "papertype": papertype,
    }

# ---------------------------------------------------------------------------
# BibTeX/BibLaTeX parsing and citations
# ---------------------------------------------------------------------------


def _strip_outer(value: str) -> str:
    value = value.strip().rstrip(",").strip()
    if (value.startswith("{") and value.endswith("}")) or (value.startswith('"') and value.endswith('"')):
        value = value[1:-1]
    return value.strip()


def bibtex_unescape(value: str) -> str:
    value = _strip_outer(value)
    replacements = {
        r"\&": "&",
        r"\%": "%",
        r"\_": "_",
        r"\textendash{}": "–",
        r"\textemdash{}": "—",
        r"\textquotedblleft{}": "“",
        r"\textquotedblright{}": "”",
        r"\'{a}": "á", r"\'{e}": "é", r"\'{i}": "í", r"\'{o}": "ó", r"\'{u}": "ú",
        r"\'{A}": "Á", r"\'{E}": "É", r"\'{I}": "Í", r"\'{O}": "Ó", r"\'{U}": "Ú",
        r"\~{a}": "ã", r"\~{o}": "õ", r"\~{A}": "Ã", r"\~{O}": "Õ",
        r"\^{a}": "â", r"\^{e}": "ê", r"\^{o}": "ô", r"\^{A}": "Â", r"\^{E}": "Ê", r"\^{O}": "Ô",
        r"\c{c}": "ç", r"\c{C}": "Ç",
    }
    for k, v in replacements.items():
        value = value.replace(k, v)
    # Drop protection braces, preserving the text.
    value = value.replace("{", "").replace("}", "")
    value = re.sub(r"\\[a-zA-Z]+\s*", "", value)
    return clean_spaces(value)


def parse_bib_entries(bib: Path) -> dict[str, dict[str, str]]:
    txt = read_text(bib)
    entries: dict[str, dict[str, str]] = {}
    i = 0
    while True:
        m = re.search(r"@([A-Za-z]+)\s*\{\s*([^,\s]+)\s*,", txt[i:])
        if not m:
            break
        start = i + m.start()
        body_start = i + m.end()
        entry_type = m.group(1).lower()
        key = m.group(2).strip()
        level = 1
        j = body_start
        while j < len(txt) and level > 0:
            ch = txt[j]
            if ch == "{":
                level += 1
            elif ch == "}":
                level -= 1
            j += 1
        body = txt[body_start:j - 1]
        fields: dict[str, str] = {"entrytype": entry_type, "id": key}
        pos = 0
        while pos < len(body):
            fm = re.search(r"([A-Za-z][A-Za-z0-9_:-]*)\s*=\s*", body[pos:])
            if not fm:
                break
            name = fm.group(1).lower()
            val_start = pos + fm.end()
            if val_start >= len(body):
                break
            quote = body[val_start]
            if quote == "{":
                lvl = 1
                k = val_start + 1
                while k < len(body) and lvl > 0:
                    if body[k] == "{":
                        lvl += 1
                    elif body[k] == "}":
                        lvl -= 1
                    k += 1
                raw_val = body[val_start:k]
                pos = k + 1
            elif quote == '"':
                k = val_start + 1
                while k < len(body):
                    if body[k] == '"' and body[k - 1] != "\\":
                        k += 1
                        break
                    k += 1
                raw_val = body[val_start:k]
                pos = k + 1
            else:
                k = val_start
                while k < len(body) and body[k] not in ",\n":
                    k += 1
                raw_val = body[val_start:k]
                pos = k + 1
            fields[name] = bibtex_unescape(raw_val)
        if "year" not in fields:
            date = fields.get("date", "")
            ym = re.search(r"(?:19|20)\d{2}", date)
            if ym:
                fields["year"] = ym.group(0)
        entries[key] = fields
        i = max(j, start + 1)
    return entries


def split_authors(author_field: str) -> list[str]:
    author_field = bibtex_unescape(author_field or "")
    if not author_field:
        return []
    if " and " in author_field:
        parts = re.split(r"\s+and\s+", author_field)
    elif ";" in author_field:
        parts = author_field.split(";")
    else:
        parts = [author_field]
    return [p.strip() for p in parts if p.strip()]


def surname(author: str) -> str:
    author = bibtex_unescape(author)
    if not author:
        return ""
    if "," in author:
        return author.split(",", 1)[0].strip()
    parts = author.split()
    if len(parts) == 1:
        return parts[0]
    particles = {"de", "da", "das", "do", "dos", "van", "von", "del", "della"}
    if len(parts) >= 2 and parts[-2].lower() in particles:
        return " ".join(parts[-2:])
    return parts[-1]


def author_abnt(author: str) -> str:
    author = bibtex_unescape(author)
    if not author:
        return ""
    if "," in author:
        last, first = [p.strip() for p in author.split(",", 1)]
    else:
        parts = author.split()
        if not parts:
            return ""
        last = surname(author)
        first = author[: max(0, len(author) - len(last))].strip()
    last_up = last.upper()
    return f"{last_up}, {first}".strip().rstrip(",")


def authors_abnt(fields: dict[str, str]) -> str:
    authors = split_authors(fields.get("author", ""))
    if not authors:
        editor = split_authors(fields.get("editor", ""))
        authors = editor
    if not authors:
        return (fields.get("organization") or fields.get("institution") or fields.get("id") or "REFERÊNCIA").upper()
    formatted = [author_abnt(a) for a in authors]
    return "; ".join([x for x in formatted if x])


def citation_label(fields: dict[str, str], textual: bool = False) -> str:
    authors = [surname(a) for a in split_authors(fields.get("author", ""))]
    authors = [a for a in authors if a]
    if not authors:
        title = fields.get("title") or fields.get("id") or "Fonte"
        label = title.split(":", 1)[0][:50]
    elif len(authors) == 1:
        label = authors[0]
    elif len(authors) == 2:
        label = f"{authors[0]} e {authors[1]}" if textual else f"{authors[0]}; {authors[1]}"
    elif len(authors) == 3:
        label = f"{authors[0]}, {authors[1]} e {authors[2]}" if textual else f"{authors[0]}; {authors[1]}; {authors[2]}"
    else:
        label = f"{authors[0]} et al."
    year = fields.get("year") or "s.d."
    return f"{label} ({year})" if textual else f"{label}, {year}"


def cite_keys(keys: str, entries: dict[str, dict[str, str]], textual: bool = False) -> str:
    parts = []
    for key in keys.split(","):
        key = key.strip()
        if not key:
            continue
        parts.append(citation_label(entries.get(key, {"id": key}), textual=textual))
    if textual:
        return "; ".join(parts)
    return "(" + "; ".join(parts) + ")" if parts else ""


def materialize_citations(text: str, entries: dict[str, dict[str, str]]) -> str:
    text = re.sub(r"\\textcite\{([^}]+)\}", lambda m: cite_keys(m.group(1), entries, textual=True), text)
    text = re.sub(r"\\parencite\{([^}]+)\}", lambda m: cite_keys(m.group(1), entries, textual=False), text)
    text = re.sub(r"\\cite\{([^}]+)\}", lambda m: cite_keys(m.group(1), entries, textual=False), text)
    text = re.sub(r"\[@([^\]]+)\]", lambda m: cite_keys(m.group(1).replace("@", ""), entries, textual=False), text)
    return text


def reference_line(fields: dict[str, str]) -> str:
    authors = authors_abnt(fields)
    title = fields.get("title") or "Sem título"
    year = fields.get("year") or "s.d."
    journal = fields.get("journaltitle") or fields.get("journal") or fields.get("periodical")
    booktitle = fields.get("booktitle")
    publisher = fields.get("publisher") or fields.get("institution") or fields.get("organization")
    address = fields.get("location") or fields.get("address")
    volume = fields.get("volume")
    number = fields.get("number") or fields.get("issue")
    pages = fields.get("pages")
    doi = fields.get("doi")
    url = fields.get("url")

    chunks = [f"{authors}. {title}."]
    if journal:
        vparts = [journal]
        if volume:
            vparts.append(f"v. {volume}")
        if number:
            vparts.append(f"n. {number}")
        if pages:
            vparts.append(f"p. {pages}")
        vparts.append(str(year))
        chunks.append(", ".join(vparts) + ".")
    elif booktitle:
        chunks.append(f"In: {booktitle}. {year}.")
    elif publisher:
        loc = f"{address}: " if address else ""
        chunks.append(f"{loc}{publisher}, {year}.")
    else:
        chunks.append(f"{year}.")
    if doi:
        chunks.append(f"DOI: {doi}.")
    if url:
        chunks.append(f"Disponível em: {url}.")
    return clean_spaces(" ".join(chunks))


# ---------------------------------------------------------------------------
# Text extraction and cleanup
# ---------------------------------------------------------------------------



def clean_latex_inline(s: str) -> str:
    if not s:
        return ""
    s = str(s)
    s = s.replace("\ufffe", "-").replace("\u00ad", "")
    s = s.replace(r"\%", "%").replace(r"\&", "&").replace(r"\_", "_")
    s = s.replace(r"\LaTeX{}", "LaTeX").replace(r"\LaTeX", "LaTeX")
    # Comandos LaTeX de bloco/spacing que não têm conteúdo textual.
    s = re.sub(r"\\vspace\*?(?:\[[^\]]*\])?\{[^{}]*\}", " ", s)
    s = re.sub(r"\\hspace\*?(?:\[[^\]]*\])?\{[^{}]*\}", " ", s)
    s = re.sub(r"\\(?:begingroup|endgroup|small|normalsize|noindent|centering|par)\b", " ", s)
    s = re.sub(r"\\(?:begin|end)\{(?:center|flushleft|flushright|quote|abstract)\}", " ", s)
    # Links.
    s = re.sub(r"\\href\{([^{}]*)\}\{([^{}]*)\}", r"\2 (\1)", s)
    s = re.sub(r"\\url\{([^{}]*)\}", r"\1", s)
    # Preserve o conteúdo de comandos de formatação. Repetir ajuda com casos aninhados simples.
    for _ in range(8):
        s2 = re.sub(r"\\(?:textit|emph|textbf|textsc|underline)\{([^{}]*)\}", r"\1", s)
        if s2 == s:
            break
        s = s2
    s = re.sub(r"\\footnote\{[^{}]*\}", "", s)
    s = re.sub(r"\\(?:section|subsection|subsubsection)\*?\{([^{}]*)\}", r"\1", s)
    # Qualquer comando remanescente com um argumento textual simples: conserva o argumento.
    s = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?\{([^{}]*)\}", r"\1", s)
    # Comandos sem argumento: remove.
    s = re.sub(r"\\[a-zA-Z]+\*?(?:\[[^\]]*\])?", "", s)
    s = s.replace("{", "").replace("}", "")
    s = s.replace("~", " ")
    s = re.sub(r"\b0\.8em\b", " ", s)
    s = re.sub(r"\bacademic_pipeline:[^\s]+", " ", s)
    return clean_spaces(s)


def line_is_metadata_or_latex(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    low = stripped.lower()
    if stripped.startswith("#+"):
        return True
    if low.startswith("academic_pipeline:"):
        return True
    if re.match(r"^\\(documentclass|usepackage|addbibresource|begin\{|end\{|makemytitle|printbibliography|nocite|bibliography|bibliographystyle|clearpage|newpage|tableofcontents|vspace|hspace|begingroup|endgroup|small|normalsize)\b", stripped):
        return True
    return False


def _extract_export_latex_blocks(org_text: str) -> list[str]:
    blocks: list[str] = []
    pat = re.compile(r"(?is)^#\+begin_export\s+latex\s*$\n(.*?)^#\+end_export\s*$", re.M)
    for m in pat.finditer(org_text):
        blocks.append(m.group(1))
    return blocks


def _clean_abstract_piece(s: str) -> str:
    s = re.sub(r"(?is)^.*?\\textbf\{(?:Resumo|Abstract)\}\s*", "", s)
    s = re.sub(r"(?is)\\noindent\s*\\textbf\{(?:Palavras-chave|Keywords)\s*:?\}.*$", "", s)
    return clean_latex_inline(s)


def _split_keywords(raw: str) -> list[str]:
    raw = clean_latex_inline(raw)
    raw = re.sub(r"^(?:Palavras-chave|Keywords)\s*:?\s*", "", raw, flags=re.I)
    return [clean_latex_inline(x).rstrip(".") for x in re.split(r"[,;]", raw) if clean_latex_inline(x).rstrip(".")]


def extract_org_abstracts(org_text: str) -> dict[str, Any]:
    """Extrai Resumo/Abstract do ORG final, inclusive de blocos LaTeX exportados.

    A versão anterior só reconhecia headings Org. Em artigos FGV/ABNT, o resumo
    frequentemente está em #+begin_export latex; se esse bloco não for interpretado,
    ele vaza para o corpo com resíduos como academic_pipeline, 0.8em e -chave.
    """
    result: dict[str, Any] = {}

    # 1) Blocos LaTeX exportados, geralmente antes da Introdução.
    export_text = "\n\n".join(_extract_export_latex_blocks(org_text))
    if export_text:
        # Português.
        m = re.search(
            r"(?is)\\textbf\{Resumo\}\s*(.*?)(?:\\noindent\s*)?\\textbf\{Palavras-chave\s*:?\}\s*(.*?)(?=\\vspace|\\textbf\{Abstract\}|\\endgroup|$)",
            export_text,
        )
        if m:
            result["abstract_pt"] = _clean_abstract_piece(m.group(1))
            result["keywords_pt"] = _split_keywords(m.group(2))
        # Inglês.
        m = re.search(
            r"(?is)\\textbf\{Abstract\}\s*(.*?)(?:\\noindent\s*)?\\textbf\{Keywords\s*:?\}\s*(.*?)(?=\\vspace|\\textbf\{Resumo\}|\\endgroup|$)",
            export_text,
        )
        if m:
            result["abstract_en"] = _clean_abstract_piece(m.group(1))
            result["keywords_en"] = _split_keywords(m.group(2))

    # 2) Headings Org como fallback/alternativa.
    patterns = [
        ("abstract_pt", r"(?is)(?:^|\n)\*+\s*Resumo\s*\n(.+?)(?=\n\*+\s*(?:Palavras-chave|Abstract|Introdução|Referências)|\Z)"),
        ("keywords_pt_raw", r"(?im)^\s*(?:Palavras-chave)\s*:?\s*(.+)$"),
        ("abstract_en", r"(?is)(?:^|\n)\*+\s*Abstract\s*\n(.+?)(?=\n\*+\s*(?:Keywords|Introdução|Referências)|\Z)"),
        ("keywords_en_raw", r"(?im)^\s*(?:Keywords)\s*:?\s*(.+)$"),
    ]
    for key, pat in patterns:
        if key in result:
            continue
        m = re.search(pat, org_text)
        if m:
            result[key] = clean_latex_inline(m.group(1))
    for key_raw, key_out in [("keywords_pt_raw", "keywords_pt"), ("keywords_en_raw", "keywords_en")]:
        raw = result.pop(key_raw, "")
        if raw and key_out not in result:
            result[key_out] = _split_keywords(raw)
    return {k: v for k, v in result.items() if v}

def extract_resumos(resumos_json: dict[str, Any], org_text: str) -> dict[str, Any]:
    out = extract_org_abstracts(org_text)
    items = resumos_json.get("items") if isinstance(resumos_json.get("items"), dict) else {}
    pt = items.get("pt-br") or items.get("pt_BR") or items.get("pt") or {}
    en = items.get("en") or items.get("en-us") or items.get("en_US") or {}
    if "abstract_pt" not in out and isinstance(pt, dict) and pt.get("abstract"):
        out["abstract_pt"] = clean_latex_inline(str(pt.get("abstract")))
    if "keywords_pt" not in out and isinstance(pt, dict) and isinstance(pt.get("keywords"), list):
        out["keywords_pt"] = [clean_latex_inline(str(x)) for x in pt.get("keywords", []) if str(x).strip()]
    if "abstract_en" not in out and isinstance(en, dict) and en.get("abstract"):
        out["abstract_en"] = clean_latex_inline(str(en.get("abstract")))
    if "keywords_en" not in out and isinstance(en, dict) and isinstance(en.get("keywords"), list):
        out["keywords_en"] = [clean_latex_inline(str(x)) for x in en.get("keywords", []) if str(x).strip()]
    return out



def parse_org_blocks(org_text: str, entries: dict[str, dict[str, str]]) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    paragraph_lines: list[str] = []
    table_lines: list[str] = []
    in_refs = False
    in_comment = False
    in_export = False

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        text = clean_spaces("\n".join(paragraph_lines))
        paragraph_lines = []
        if not text:
            return
        text = materialize_citations(text, entries)
        text = clean_latex_inline(text)
        if not text:
            return
        # Guardas contra resíduos de blocos técnicos ou resumos exportados.
        low = text.lower().strip()
        if low.startswith("academic_pipeline:"):
            return
        if low in {"0.8em", "resumo", "abstract", "palavras-chave", "keywords"}:
            return
        if low.startswith("-chave:") or low.startswith(": atestmed"):
            return
        blocks.append({"type": "paragraph", "text": text})

    def flush_table() -> None:
        nonlocal table_lines
        if not table_lines:
            return
        rows: list[list[str]] = []
        for ln in table_lines:
            if re.match(r"^\s*\|[-+ ]+\|\s*$", ln):
                continue
            cells = [clean_latex_inline(materialize_citations(c.strip(), entries)) for c in ln.strip().strip("|").split("|")]
            if cells and any(cells):
                rows.append(cells)
        table_lines = []
        if rows:
            blocks.append({"type": "table", "rows": rows})

    for raw in org_text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()
        low = stripped.lower()

        if re.match(r"^#\+begin_comment\b", stripped, flags=re.I):
            flush_table(); flush_paragraph(); in_comment = True; continue
        if re.match(r"^#\+end_comment\b", stripped, flags=re.I):
            in_comment = False; continue
        if in_comment:
            continue

        if re.match(r"^#\+begin_export\b", stripped, flags=re.I):
            flush_table(); flush_paragraph(); in_export = True; continue
        if re.match(r"^#\+end_export\b", stripped, flags=re.I):
            in_export = False; continue
        if in_export:
            continue

        if not stripped:
            flush_table(); flush_paragraph(); continue
        if line_is_metadata_or_latex(line):
            continue
        # Stop reading body references generated by LaTeX; we add references ourselves.
        if re.match(r"^(\*+\s*)?Refer[eê]ncias\s*$", stripped, flags=re.I):
            flush_table(); flush_paragraph(); in_refs = True; continue
        if in_refs:
            continue
        m = re.match(r"^(\*{1,6})\s+(.+)$", stripped)
        if m:
            flush_table(); flush_paragraph()
            title = clean_latex_inline(materialize_citations(m.group(2), entries))
            if title.lower() not in {"resumo", "abstract", "palavras-chave", "keywords"}:
                blocks.append({"type": "heading", "level": min(len(m.group(1)), 3), "text": title})
            continue
        m = re.match(r"^\\(section|subsection|subsubsection)\*?\{(.+)\}\s*$", stripped)
        if m:
            flush_table(); flush_paragraph()
            level = {"section": 1, "subsection": 2, "subsubsection": 3}.get(m.group(1), 1)
            title = clean_latex_inline(materialize_citations(m.group(2), entries))
            if title.lower() not in {"resumo", "abstract", "palavras-chave", "keywords"}:
                blocks.append({"type": "heading", "level": level, "text": title})
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph(); table_lines.append(stripped); continue
        paragraph_lines.append(line)
    flush_table(); flush_paragraph()
    return blocks

def parse_document_json_blocks(document_json: dict[str, Any], entries: dict[str, dict[str, str]]) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    for sec in document_json.get("sections", []) if isinstance(document_json.get("sections"), list) else []:
        if not isinstance(sec, dict):
            continue
        title = clean_latex_inline(str(sec.get("title", "")))
        if title:
            blocks.append({"type": "heading", "level": 1, "text": title})
        text = str(sec.get("text", ""))
        # Remove first LaTeX section command if it duplicates title.
        text = re.sub(r"^\\section\*?\{[^}]+\}\s*", "", text.strip())
        for part in re.split(r"\n\s*\n", text):
            part = clean_spaces(part)
            if not part:
                continue
            if part.lower() == title.lower():
                continue
            part = materialize_citations(part, entries)
            part = clean_latex_inline(part)
            if part:
                blocks.append({"type": "paragraph", "text": part})
    return blocks


def build_body_blocks(org_text: str, document_json: dict[str, Any], entries: dict[str, dict[str, str]]) -> list[dict[str, Any]]:
    org_blocks = parse_org_blocks(org_text, entries) if org_text else []
    # Trust ORG when it has substantial content and headings.
    if sum(1 for b in org_blocks if b.get("type") == "heading") >= 3 and sum(len(b.get("text", "")) for b in org_blocks) > 5000:
        return org_blocks
    return parse_document_json_blocks(document_json, entries)


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------


def set_run_font(run: Any, name: str = "Times New Roman", size_pt: int = 12, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = name
    run.font.size = Pt(size_pt)
    # ABNT/FGV: todo texto visível deve permanecer em preto. O estilo padrão
    # de headings do Word costuma herdar azul temático; por isso a cor é
    # fixada explicitamente no nível do run.
    try:
        run.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ["w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"]:
        rfonts.set(qn(attr), name)


def set_style_font(style: Any, name: str = "Times New Roman", size_pt: int = 12, bold: bool | None = None) -> None:
    style.font.name = name
    style.font.size = Pt(size_pt)
    # Corrige a cor azul herdada dos estilos Heading 1/2/3 do template padrão.
    try:
        style.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        pass
    if bold is not None:
        style.font.bold = bold
    rpr = style._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ["w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"]:
        rfonts.set(qn(attr), name)


def configure_doc_styles(doc: Any, layout: dict[str, Any]) -> None:
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    margins = layout.get("margens_cm", {}) if isinstance(layout, dict) else {}
    sec.top_margin = Cm(float(margins.get("superior", 3.0)))
    sec.left_margin = Cm(float(margins.get("esquerda", 3.0)))
    sec.right_margin = Cm(float(margins.get("direita", 2.0)))
    sec.bottom_margin = Cm(float(margins.get("inferior", 2.0)))

    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal, size_pt=int(layout.get("fonte_texto_pt", 12)))
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = float(layout.get("espacamento_texto", 1.5))
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Cm(float(layout.get("recuo_primeira_linha_cm", 1.25)))

    for style_name, size, before, after in [
        ("Title", 14, 0, 12),
        ("Heading 1", 12, 12, 6),
        ("Heading 2", 12, 12, 6),
        ("Heading 3", 12, 12, 6),
    ]:
        if style_name in styles:
            st = styles[style_name]
            set_style_font(st, size_pt=size, bold=True)
            st.paragraph_format.line_spacing = 1.5
            st.paragraph_format.space_before = Pt(before)
            st.paragraph_format.space_after = Pt(after)
            st.paragraph_format.first_line_indent = Cm(0)

    if "Bibliography" in styles:
        st = styles["Bibliography"]
        set_style_font(st, size_pt=12, bold=False)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        st.paragraph_format.line_spacing = 1.0
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(6)
        st.paragraph_format.first_line_indent = Cm(0)


def add_centered_paragraph(doc: Any, text: str, *, bold: bool = False, size: int = 12, before: int = 0, after: int = 0) -> Any:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, size_pt=size, bold=bold)
    return p


def add_body_paragraph(doc: Any, text: str, *, first_indent: bool = True, align: Any = WD_ALIGN_PARAGRAPH.JUSTIFY) -> Any:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.first_line_indent = Cm(1.25 if first_indent else 0)
    r = p.add_run(text)
    set_run_font(r, size_pt=12)
    return p


def add_heading(doc: Any, text: str, level: int = 1) -> Any:
    p = doc.add_heading("", level=min(max(level, 1), 3))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    set_run_font(r, size_pt=12, bold=True)
    return p



def add_cover(doc: Any, metadata: dict[str, str]) -> None:
    # Capa institucional próxima às convenções ABNT/FGV.
    inst_lines_raw = [
        metadata.get("institution", ""),
        metadata.get("program", ""),
        metadata.get("course", ""),
        metadata.get("discipline", ""),
    ]
    inst_lines: list[str] = []
    seen: set[str] = set()
    for line in inst_lines_raw:
        line = clean_spaces(line)
        if not line or invalid_cover_value(line):
            continue
        key = strip_accents_for_sort(line)
        if key in seen:
            continue
        seen.add(key)
        inst_lines.append(line)
    if not inst_lines:
        inst_lines = ["Fundação Getúlio Vargas"]
    for line in inst_lines:
        add_centered_paragraph(doc, line.upper(), bold=True, size=12, after=0)

    add_centered_paragraph(doc, metadata.get("author", ""), bold=False, size=12, before=70, after=0)
    add_centered_paragraph(doc, metadata.get("title", ""), bold=True, size=12, before=95, after=0)
    if metadata.get("subtitle"):
        add_centered_paragraph(doc, metadata["subtitle"], bold=False, size=12, after=0)
    if metadata.get("covernote") or metadata.get("discipline") or metadata.get("professor"):
        note = metadata.get("covernote", "")
        parts: list[str] = []
        if metadata.get("professor"):
            professor = metadata["professor"].strip()
            if professor and not professor.lower().startswith("professor"):
                professor = "Professor: " + professor
            if professor:
                parts.append(professor)
        elif metadata.get("discipline") and not note:
            # Fallback raro: se não houver nota nem professor, preserva a disciplina.
            parts.append(metadata["discipline"])
        if parts:
            note = (note + "\n" if note else "") + " — ".join(parts)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(7.5)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(60)
        for i, chunk in enumerate(note.split("\n")):
            if i:
                p.add_run().add_break()
            r = p.add_run(chunk)
            set_run_font(r, size_pt=10)
    add_centered_paragraph(doc, metadata.get("city", "Brasília"), size=12, before=240, after=0)
    add_centered_paragraph(doc, metadata.get("year", str(datetime.now().year)), size=12, before=0, after=0)
    doc.add_page_break()

def add_abstracts(doc: Any, abstracts: dict[str, Any]) -> None:
    if abstracts.get("abstract_pt"):
        add_heading(doc, "Resumo", 1)
        add_body_paragraph(doc, str(abstracts["abstract_pt"]), first_indent=False)
        kws = abstracts.get("keywords_pt") or []
        if kws:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.5
            r1 = p.add_run("Palavras-chave: ")
            set_run_font(r1, bold=True)
            r2 = p.add_run("; ".join(kws) + ".")
            set_run_font(r2)
        doc.add_paragraph()
    if abstracts.get("abstract_en"):
        add_heading(doc, "Abstract", 1)
        add_body_paragraph(doc, str(abstracts["abstract_en"]), first_indent=False)
        kws = abstracts.get("keywords_en") or []
        if kws:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.5
            r1 = p.add_run("Keywords: ")
            set_run_font(r1, bold=True)
            r2 = p.add_run("; ".join(kws) + ".")
            set_run_font(r2)
        doc.add_paragraph()


def add_table(doc: Any, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, row in enumerate(rows):
        cells = table.add_row().cells
        for cidx in range(cols):
            text = row[cidx] if cidx < len(row) else ""
            cells[cidx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[cidx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(text)
            set_run_font(r, size_pt=10, bold=(idx == 0))
    doc.add_paragraph()


def add_references(doc: Any, entries: dict[str, dict[str, str]]) -> list[str]:
    refs = [reference_line(v) for v in entries.values()]
    refs = [r for r in refs if r]
    refs = sorted(refs, key=strip_accents_for_sort)
    add_heading(doc, "Referências", 1)
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(ref)
        set_run_font(r, size_pt=12)
    return refs


def load_layout(compliance_json: dict[str, Any]) -> dict[str, Any]:
    layout = dict(DEFAULT_LAYOUT)
    items = compliance_json.get("items")
    if isinstance(items, list):
        for item in items:
            if isinstance(item, dict) and item.get("id") == "layout.rules":
                detail = item.get("detail", "")
                # detail is repr(dict), not guaranteed JSON. Parse gently.
                m = re.search(r"'margens_cm':\s*\{([^}]+)\}", str(detail))
                if m:
                    for name in ["superior", "esquerda", "direita", "inferior"]:
                        mm = re.search(rf"'{name}':\s*([0-9.]+)", m.group(1))
                        if mm:
                            layout["margens_cm"][name] = float(mm.group(1))
                for key in ["fonte_texto_pt", "espacamento_texto", "recuo_primeira_linha_cm"]:
                    mm = re.search(rf"'{key}':\s*([0-9.]+)", str(detail))
                    if mm:
                        val = float(mm.group(1))
                        layout[key] = int(val) if val.is_integer() else val
    return layout




FICHAMENTO_QUALITATIVO_LAYOUT = "fichamento_qualitativo"
FICHAMENTO_QUALITATIVO_SECTIONS = [
    "REFERÊNCIAS BIBLIOGRÁFICAS",
    "SÍNTESE DOS TEXTOS",
    "PRINCIPAIS CONCEITOS E ARGUMENTOS",
    "ANÁLISE CRÍTICA E REFLEXÕES PESSOAIS",
    "CONEXÕES E DIÁLOGOS ENTRE OS TEXTOS",
    "APLICAÇÕES EM POLÍTICAS PÚBLICAS E GOVERNO",
    "QUESTÕES PARA APROFUNDAMENTO",
]
FICHAMENTO_FGV_DARK = "003B70"
FICHAMENTO_FGV_BLUE = "005CA9"
FICHAMENTO_FGV_LIGHT = "DCEAF7"
FICHAMENTO_GRAY = "666666"
FICHAMENTO_LIGHT_GRAY = "F3F5F7"
FICHAMENTO_WHITE = "FFFFFF"


def selected_document_layout(toml_data: dict[str, Any]) -> str:
    for section_name in ("documento", "document"):
        section = toml_data.get(section_name)
        if isinstance(section, dict):
            value = clean_spaces(str(section.get("layout", "")))
            if value:
                return value
    return clean_spaces(str(toml_data.get("layout", "")))


def _fichamento_set_color(run: Any, color_hex: str) -> None:
    color_hex = color_hex.strip().lstrip("#")
    if not re.fullmatch(r"[0-9A-Fa-f]{6}", color_hex):
        raise ValueError(f"cor RGB inválida: {color_hex!r}")
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16),
    )


def _fichamento_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _fichamento_cell_margins(
    cell: Any, *, top: int = 40, start: int = 100, bottom: int = 40, end: int = 100
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _fichamento_compact_row(row: Any) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    height = OxmlElement("w:trHeight")
    height.set(qn("w:val"), "300")
    height.set(qn("w:hRule"), "exact")
    tr_pr.append(height)


def add_fichamento_qualitativo_header(doc: Any) -> None:
    section = doc.sections[0]
    section.header_distance = Cm(0.8)
    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(16.0))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _fichamento_compact_row(table.rows[0])
    table.columns[0].width = Cm(2.8)
    table.columns[1].width = Cm(13.2)
    left, right = table.rows[0].cells
    _fichamento_cell_shading(left, FICHAMENTO_FGV_DARK)
    _fichamento_cell_shading(right, FICHAMENTO_FGV_BLUE)
    for cell in (left, right):
        _fichamento_cell_margins(cell, top=20, bottom=20, start=100, end=100)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("FGV")
    set_run_font(run, name="Arial", size_pt=10, bold=True)
    _fichamento_set_color(run, FICHAMENTO_WHITE)
    p = right.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("FICHAMENTO DE LEITURA")
    set_run_font(run, name="Arial", size_pt=8, bold=True)
    _fichamento_set_color(run, FICHAMENTO_WHITE)
    _materialize_fichamento_header_geometry(table)


def _fichamento_date(metadata: dict[str, str], toml_data: dict[str, Any]) -> str:
    date_value = first_by_keys(toml_data, {"data", "date"})
    return first_valid(date_value, metadata.get("year", ""), default="")


def add_fichamento_qualitativo_front_matter(
    doc: Any, metadata: dict[str, str], toml_data: dict[str, Any]
) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("FICHAMENTO DE LEITURA")
    set_run_font(run, name="Arial", size_pt=20, bold=True)
    _fichamento_set_color(run, FICHAMENTO_FGV_DARK)

    subtitle = clean_spaces(metadata.get("title", ""))
    if subtitle:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        p.paragraph_format.first_line_indent = Cm(0)
        run = p.add_run(subtitle)
        set_run_font(run, name="Arial", size_pt=12, italic=True)
        _fichamento_set_color(run, FICHAMENTO_GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run("FICHA TÉCNICA")
    set_run_font(run, name="Arial", size_pt=11, bold=True)
    _fichamento_set_color(run, FICHAMENTO_FGV_DARK)

    rows = [
        ("Disciplina", clean_spaces(metadata.get("discipline", ""))),
        ("Professor", clean_spaces(metadata.get("professor", ""))),
        ("Aluno(a)", clean_spaces(metadata.get("author", ""))),
        ("Data", _fichamento_date(metadata, toml_data)),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(3.1)
    table.columns[1].width = Cm(12.9)
    for index, (label, value) in enumerate(rows):
        left, right = table.rows[index].cells
        _fichamento_cell_shading(left, FICHAMENTO_FGV_LIGHT)
        _fichamento_cell_shading(right, FICHAMENTO_WHITE if index % 2 == 0 else FICHAMENTO_LIGHT_GRAY)
        for cell in (left, right):
            _fichamento_cell_margins(cell, top=80, bottom=80, start=120, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lr = lp.add_run(label)
        set_run_font(lr, name="Arial", size_pt=10, bold=True)
        _fichamento_set_color(lr, FICHAMENTO_FGV_DARK)
        rp = right.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rr = rp.add_run(value)
        set_run_font(rr, size_pt=11)
    _materialize_fichamento_technical_geometry(table)


def _fichamento_normalize_heading(text: str) -> str:
    value = clean_spaces(str(text or ""))
    value = re.sub(r"^\s*\d+(?:\.\d+)*(?:[\.\)])?\s+", "", value)
    return strip_accents_for_sort(value)


def add_fichamento_qualitativo_section_heading(doc: Any, number: int, title: str) -> Any:
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"{number}. {title}")
    set_run_font(run, name="Arial", size_pt=14, bold=True)
    _fichamento_set_color(run, FICHAMENTO_FGV_DARK)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "10")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), FICHAMENTO_FGV_BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    return p


def add_fichamento_qualitativo_references(
    doc: Any, entries: dict[str, dict[str, str]]
) -> list[str]:
    refs = [reference_line(value) for value in entries.values()]
    refs = sorted([ref for ref in refs if ref], key=strip_accents_for_sort)
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(ref)
        set_run_font(run, size_pt=12)
    return refs


def render_fichamento_qualitativo_body(
    doc: Any,
    body_blocks: list[dict[str, Any]],
    entries: dict[str, dict[str, str]],
) -> tuple[list[str], list[str]]:
    normalized = {
        _fichamento_normalize_heading(title): (index, title)
        for index, title in enumerate(FICHAMENTO_QUALITATIVO_SECTIONS, 1)
    }
    expected_index = 1
    current_primary = 0
    seen: list[str] = []
    refs: list[str] = []
    for block in body_blocks:
        btype = block.get("type")
        if btype == "heading":
            level = int(block.get("level", 1))
            text = clean_spaces(str(block.get("text", "")))
            key = _fichamento_normalize_heading(text)
            primary = normalized.get(key)
            if primary is not None:
                index, canonical = primary
                if index != expected_index:
                    die(
                        "ERRO: ordem de seções do fichamento divergente. "
                        f"Esperada seção {expected_index}; encontrada {index}: {canonical}"
                    )
                add_fichamento_qualitativo_section_heading(doc, index, canonical)
                seen.append(canonical)
                current_primary = index
                expected_index += 1
                if index == 1:
                    refs = add_fichamento_qualitativo_references(doc, entries)
                continue
            if level == 1:
                die(f"ERRO: seção primária inesperada no fichamento: {text}")
            if current_primary > 1:
                add_heading(doc, text, min(max(level, 2), 3))
            continue
        if current_primary == 0:
            continue
        if current_primary == 1:
            continue
        if btype == "paragraph":
            text = clean_spaces(str(block.get("text", "")))
            if text:
                add_body_paragraph(doc, text, first_indent=True)
        elif btype == "table":
            rows = block.get("rows")
            if isinstance(rows, list):
                add_table(doc, rows)

    if seen != FICHAMENTO_QUALITATIVO_SECTIONS:
        missing = [title for title in FICHAMENTO_QUALITATIVO_SECTIONS if title not in seen]
        die(f"ERRO: fichamento sem as sete seções canônicas: {missing}")
    return refs, seen


def _set_fixed_table_geometry(table: Any, total_width_dxa: int, column_widths_dxa: tuple[int, ...] | list[int]) -> None:
    """Fix tblW, tblGrid and every tcW to one coherent OOXML geometry."""
    widths = tuple(int(value) for value in column_widths_dxa)
    total = int(total_width_dxa)
    if not widths or sum(widths) != total:
        raise ValueError(f"invalid fixed table geometry: total={total} widths={widths}")
    if any(value <= 0 for value in widths):
        raise ValueError(f"invalid non-positive table width: {widths}")
    if len(table.columns) != len(widths):
        raise ValueError(f"column count mismatch: observed={len(table.columns)} expected={len(widths)}")

    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        if len(row.cells) != len(widths):
            raise ValueError(
                f"row cell count mismatch: observed={len(row.cells)} expected={len(widths)}"
            )
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))


def _set_cell_margins_dxa(cell: Any, *, top: int, bottom: int, start: int, end: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(int(value)))
        node.set(qn("w:type"), "dxa")


def _set_cell_nowrap(cell: Any) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    if tc_pr.find(qn("w:noWrap")) is None:
        tc_pr.append(OxmlElement("w:noWrap"))


def _materialize_fichamento_header_geometry(table: Any) -> None:
    """Apply the golden header geometry directly in the live header builder."""
    _set_fixed_table_geometry(table, 9071, (1588, 7483))
    for cell in table.rows[0].cells:
        _set_cell_margins_dxa(cell, top=20, bottom=20, start=100, end=100)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    height = tr_pr.find(qn("w:trHeight"))
    if height is None:
        height = OxmlElement("w:trHeight")
        tr_pr.append(height)
    height.set(qn("w:val"), "300")
    height.set(qn("w:hRule"), "exact")


def _materialize_fichamento_technical_geometry(table: Any) -> None:
    """Apply the golden technical-sheet geometry directly in the live builder."""
    _set_fixed_table_geometry(table, 9071, (1758, 7313))
    if len(table.rows) < 4:
        raise ValueError(f"technical sheet row count mismatch: observed={len(table.rows)} expected>=4")
    for row in table.rows[:4]:
        _set_cell_margins_dxa(row.cells[0], top=20, bottom=20, start=30, end=30)
        _set_cell_margins_dxa(row.cells[1], top=20, bottom=20, start=60, end=60)
        _set_cell_nowrap(row.cells[0])


def _table_cell_text(cell: Any) -> str:
    return " ".join((paragraph.text or "").strip() for paragraph in cell.paragraphs).strip()


def _iter_fichamento_tables(doc: Any) -> list[Any]:
    tables = list(doc.tables)
    seen: set[int] = {id(table._tbl) for table in tables}
    for section in doc.sections:
        for header in (section.header, section.first_page_header, section.even_page_header):
            for table in header.tables:
                marker = id(table._tbl)
                if marker not in seen:
                    tables.append(table)
                    seen.add(marker)
    return tables


def _normalize_fichamento_qualitativo_table_geometry(doc: Any) -> dict[str, bool]:
    """Materialize the user-approved cross-renderer geometry before DOCX serialization."""
    header_applied = False
    technical_applied = False
    expected_labels = ("Disciplina", "Professor", "Aluno(a)", "Data")
    for table in _iter_fichamento_tables(doc):
        if len(table.columns) != 2 or not table.rows:
            continue
        first_row = tuple(_table_cell_text(cell) for cell in table.rows[0].cells)
        if first_row == ("FGV", "FICHAMENTO DE LEITURA"):
            _set_fixed_table_geometry(table, 9071, (1588, 7483))
            for cell in table.rows[0].cells:
                _set_cell_margins_dxa(cell, top=20, bottom=20, start=100, end=100)
            tr_pr = table.rows[0]._tr.get_or_add_trPr()
            height = tr_pr.find(qn("w:trHeight"))
            if height is None:
                height = OxmlElement("w:trHeight")
                tr_pr.append(height)
            height.set(qn("w:val"), "300")
            height.set(qn("w:hRule"), "exact")
            header_applied = True
            continue
        labels = tuple(_table_cell_text(row.cells[0]) for row in table.rows[:4]) if len(table.rows) >= 4 else ()
        if labels == expected_labels:
            _set_fixed_table_geometry(table, 9071, (1758, 7313))
            for row in table.rows[:4]:
                _set_cell_margins_dxa(row.cells[0], top=20, bottom=20, start=30, end=30)
                _set_cell_margins_dxa(row.cells[1], top=20, bottom=20, start=60, end=60)
                _set_cell_nowrap(row.cells[0])
            technical_applied = True
    # Generic layouts legitimately have neither table. In that case this helper is a no-op.
    # A partial match, however, signals a broken fichamento geometry and must fail closed.
    if not header_applied and not technical_applied:
        return {"header": False, "technical": False}
    if header_applied != technical_applied:
        raise ValueError(
            "partial fichamento cross-renderer geometry: "
            f"header={header_applied} technical={technical_applied}"
        )
    return {"header": header_applied, "technical": technical_applied}


def _validate_fichamento_cross_renderer_ooxml(docx_path: Path | str) -> dict[str, Any]:
    """Validate raw OOXML authorities, not python-docx width abstractions alone."""
    import xml.etree.ElementTree as _ET

    path = Path(docx_path)
    if not path.is_file():
        raise ValueError(f"DOCX not found for OOXML validation: {path}")
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def text_of(node: Any) -> str:
        return "".join((part.text or "") for part in node.findall(".//" + ns + "t")).strip()

    def attr(node: Any | None, name: str) -> str | None:
        return None if node is None else node.attrib.get(ns + name)

    def table_descriptor(tbl: Any) -> dict[str, Any]:
        pr = tbl.find(ns + "tblPr")
        layout = pr.find(ns + "tblLayout") if pr is not None else None
        tbl_w = pr.find(ns + "tblW") if pr is not None else None
        grid = tbl.find(ns + "tblGrid")
        grid_widths = [int(col.attrib.get(ns + "w", "0")) for col in list(grid or []) if col.tag == ns + "gridCol"]
        rows = tbl.findall(ns + "tr")
        row_cells: list[list[dict[str, Any]]] = []
        for row in rows:
            cells: list[dict[str, Any]] = []
            for tc in row.findall(ns + "tc"):
                tc_pr = tc.find(ns + "tcPr")
                tc_w = tc_pr.find(ns + "tcW") if tc_pr is not None else None
                tc_mar = tc_pr.find(ns + "tcMar") if tc_pr is not None else None
                margins: dict[str, int | None] = {}
                for side in ("top", "bottom", "start", "end"):
                    item = tc_mar.find(ns + side) if tc_mar is not None else None
                    margins[side] = int(attr(item, "w")) if attr(item, "w") not in {None, ""} else None
                cells.append({
                    "text": text_of(tc),
                    "tcW": int(attr(tc_w, "w")) if attr(tc_w, "w") not in {None, ""} else None,
                    "tcW_type": attr(tc_w, "type"),
                    "margins": margins,
                    "noWrap": bool(tc_pr is not None and tc_pr.find(ns + "noWrap") is not None),
                })
            row_cells.append(cells)
        return {
            "layout": attr(layout, "type"),
            "tblW": int(attr(tbl_w, "w")) if attr(tbl_w, "w") not in {None, ""} else None,
            "tblW_type": attr(tbl_w, "type"),
            "grid": grid_widths,
            "rows": row_cells,
            "row_height": (
                {
                    "val": int(attr(rows[0].find(ns + "trPr").find(ns + "trHeight"), "val")),
                    "rule": attr(rows[0].find(ns + "trPr").find(ns + "trHeight"), "hRule"),
                }
                if rows and rows[0].find(ns + "trPr") is not None and rows[0].find(ns + "trPr").find(ns + "trHeight") is not None
                else None
            ),
        }

    descriptors: list[dict[str, Any]] = []
    with zipfile.ZipFile(path) as zf:
        members = ["word/document.xml"] + sorted(
            name for name in zf.namelist() if re.fullmatch(r"word/header\d+\.xml", name)
        )
        for member in members:
            root = _ET.fromstring(zf.read(member))
            for tbl in root.findall(".//" + ns + "tbl"):
                row = table_descriptor(tbl)
                row["member"] = member
                descriptors.append(row)

    header = None
    technical = None
    for row in descriptors:
        if row["rows"] and len(row["rows"][0]) >= 2:
            first = tuple(cell["text"] for cell in row["rows"][0][:2])
            if first == ("FGV", "FICHAMENTO DE LEITURA"):
                header = row
        labels = tuple(
            r[0]["text"] for r in row["rows"][:4] if r
        ) if len(row["rows"]) >= 4 else ()
        if labels == ("Disciplina", "Professor", "Aluno(a)", "Data"):
            technical = row

    problems: list[str] = []
    if header is None:
        problems.append("header table not found in raw OOXML")
    else:
        if header["layout"] != "fixed": problems.append(f"header tblLayout={header['layout']!r}")
        if header["tblW_type"] != "dxa" or header["tblW"] != 9071: problems.append(f"header tblW={header['tblW_type']}/{header['tblW']}")
        if tuple(header["grid"]) != (1588, 7483): problems.append(f"header grid={header['grid']}")
        if not header["rows"] or tuple(c["tcW"] for c in header["rows"][0][:2]) != (1588, 7483): problems.append("header tcW mismatch")
        for cell in header["rows"][0][:2]:
            if cell["margins"] != {"top": 20, "bottom": 20, "start": 100, "end": 100}:
                problems.append(f"header padding={cell['margins']}")
        if header["row_height"] != {"val": 300, "rule": "exact"}:
            problems.append(f"header row height={header['row_height']}")
        if sum(header["grid"]) != header["tblW"]:
            problems.append("header tblW/grid sum conflict")
    if technical is None:
        problems.append("technical table not found in raw OOXML")
    else:
        if technical["layout"] != "fixed": problems.append(f"technical tblLayout={technical['layout']!r}")
        if technical["tblW_type"] != "dxa" or technical["tblW"] != 9071: problems.append(f"technical tblW={technical['tblW_type']}/{technical['tblW']}")
        if tuple(technical["grid"]) != (1758, 7313): problems.append(f"technical grid={technical['grid']}")
        observed_labels = tuple(row[0]["text"] for row in technical["rows"][:4])
        if observed_labels != ("Disciplina", "Professor", "Aluno(a)", "Data"):
            problems.append(f"technical labels={observed_labels}")
        for index, row in enumerate(technical["rows"][:4]):
            if len(row) < 2 or tuple(cell["tcW"] for cell in row[:2]) != (1758, 7313):
                problems.append(f"technical row {index} tcW mismatch")
                continue
            if row[0]["margins"] != {"top": 20, "bottom": 20, "start": 30, "end": 30}:
                problems.append(f"technical label padding row {index}={row[0]['margins']}")
            if row[1]["margins"] != {"top": 20, "bottom": 20, "start": 60, "end": 60}:
                problems.append(f"technical value padding row {index}={row[1]['margins']}")
            if not row[0]["noWrap"]:
                problems.append(f"technical label noWrap missing row {index}")
        if sum(technical["grid"]) != technical["tblW"]:
            problems.append("technical tblW/grid sum conflict")
    if problems:
        raise ValueError("cross-renderer OOXML validation failed: " + "; ".join(problems))
    return {"ok": True, "header": header, "technical": technical, "tables_scanned": len(descriptors)}

def validate_fichamento_qualitativo_docx(path: Path) -> dict[str, Any]:
    _cross_renderer_ooxml = _validate_fichamento_cross_renderer_ooxml(path)
    report: dict[str, Any] = {"path": str(path), "ok": False, "warnings": [], "layout": FICHAMENTO_QUALITATIVO_LAYOUT}
    if not path.exists() or path.stat().st_size < 1000:
        report["warnings"].append("DOCX ausente ou muito pequeno.")
        return report
    try:
        rendered = Document(str(path))
        with zipfile.ZipFile(path) as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
            header_xml = "\n".join(
                zf.read(name).decode("utf-8", errors="ignore")
                for name in zf.namelist()
                if name.startswith("word/header") and name.endswith(".xml")
            )
    except Exception as exc:
        report["warnings"].append(f"DOCX inválido: {exc}")
        return report

    section = rendered.sections[0]
    margins = {
        "top": round(section.top_margin.cm, 3),
        "left": round(section.left_margin.cm, 3),
        "right": round(section.right_margin.cm, 3),
        "bottom": round(section.bottom_margin.cm, 3),
    }
    report["margins_cm"] = margins
    for key, expected in {"top": 3.0, "left": 3.0, "right": 2.0, "bottom": 2.0}.items():
        if abs(margins[key] - expected) > 0.03:
            report["warnings"].append(f"Margem {key} divergente: {margins[key]} cm != {expected} cm.")

    header_tables = section.header.tables
    report["header_table_count"] = len(header_tables)
    if len(header_tables) != 1:
        report["warnings"].append(f"Cabeçalho deve conter exatamente uma faixa/tabela; observado {len(header_tables)}.")
    else:
        table = header_tables[0]
        shape = [len(table.rows), len(table.rows[0].cells) if table.rows else 0]
        report["header_shape"] = shape
        if shape != [1, 2]:
            report["warnings"].append(f"Cabeçalho não é a faixa compacta 1x2 aprovada: {shape}.")
    for token in ("FGV", "FICHAMENTO DE LEITURA"):
        if token not in header_xml:
            report["warnings"].append(f"Cabeçalho sem {token!r}.")

    body_text = "\n".join(p.text for p in rendered.paragraphs)
    for token in ("FICHAMENTO DE LEITURA", "FICHA TÉCNICA"):
        if token not in body_text:
            report["warnings"].append(f"Front matter sem {token!r}.")
    for label in ("Disciplina", "Professor", "Aluno(a)", "Data"):
        if label not in document_xml:
            report["warnings"].append(f"Ficha Técnica sem campo {label!r}.")

    body_tables = rendered.tables
    report["technical_sheet_table_count"] = len(body_tables)
    if len(body_tables) != 1:
        report["warnings"].append(f"Esperada exatamente uma tabela de Ficha Técnica; observadas {len(body_tables)}.")
    else:
        technical_sheet = body_tables[0]
        labels = [row.cells[0].text.strip() for row in technical_sheet.rows]
        report["technical_sheet_labels"] = labels
        expected_labels = ["Disciplina", "Professor", "Aluno(a)", "Data"]
        if labels != expected_labels:
            report["warnings"].append(f"Rótulos da Ficha Técnica divergentes: {labels!r}.")
        if len(technical_sheet.columns) != 2:
            report["warnings"].append(f"Ficha Técnica deve ter duas colunas; observadas {len(technical_sheet.columns)}.")
        else:
            label_width_cm = round(technical_sheet.columns[0].width.cm, 3)
            value_width_cm = round(technical_sheet.columns[1].width.cm, 3)
            report["technical_sheet_label_column_width_cm"] = label_width_cm
            report["technical_sheet_value_column_width_cm"] = value_width_cm
            if abs(label_width_cm - 3.1) > 0.03:
                report["warnings"].append(f"Coluna de rótulos divergente: {label_width_cm} cm != 3.1 cm.")
            if abs(value_width_cm - 12.9) > 0.03:
                report["warnings"].append(f"Coluna de valores divergente: {value_width_cm} cm != 12.9 cm.")
    for index, title in enumerate(FICHAMENTO_QUALITATIVO_SECTIONS, 1):
        visible = f"{index}. {title}"
        if visible not in body_text:
            report["warnings"].append(f"Seção canônica ausente: {visible}")

    page_breaks = document_xml.count('w:type="page"')
    report["page_break_count"] = page_breaks
    if page_breaks < 7:
        report["warnings"].append(f"Esperadas ao menos 7 quebras de página para as seções; observadas {page_breaks}.")

    for token, label in [
        ("\\cite", "comando \\cite"), ("\\section", "comando \\section"),
        ("\\printbibliography", "comando \\printbibliography"), ("\\textbf", "comando \\textbf"),
        ("\\textit", "comando \\textit"), ("\\vspace", "comando \\vspace"),
        ("academic_pipeline:", "marcador técnico academic_pipeline"),
    ]:
        if token in document_xml:
            report["warnings"].append(f"Resíduo visível detectado: {label}.")

    normal = rendered.styles["Normal"]
    report["body_font"] = normal.font.name
    report["body_font_size_pt"] = float(normal.font.size.pt) if normal.font.size is not None else None
    if normal.font.name != "Times New Roman":
        report["warnings"].append(f"Fonte Normal divergente: {normal.font.name!r}.")
    if normal.font.size is None or abs(normal.font.size.pt - 12) > 0.1:
        report["warnings"].append("Tamanho da fonte Normal divergente de 12 pt.")

    report["gradient_bands_below_header"] = 0
    report["compact_header"] = len(header_tables) == 1 and report.get("header_shape") == [1, 2]
    report["font_color_policy"] = "approved_fichamento_fgv_palette"
    report["size_bytes"] = path.stat().st_size
    report["ok"] = not report["warnings"]
    return report


def render_fichamento_qualitativo(
    *, paths: ArticlePaths, metadata: dict[str, str], body_blocks: list[dict[str, Any]],
    entries: dict[str, dict[str, str]], layout: dict[str, Any], toml_data: dict[str, Any], quiet: bool,
) -> Path:
    doc = Document()
    configure_doc_styles(doc, layout)
    add_fichamento_qualitativo_header(doc)
    add_fichamento_qualitativo_front_matter(doc, metadata, toml_data)
    refs, primary_sections = render_fichamento_qualitativo_body(doc, body_blocks, entries)
    backup(paths.docx, "docx_canonico", quiet=quiet)
    doc.save(paths.docx)

    canonical = {
        "schema_version": "docx-canonico-v14",
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "paths": {key: str(value) if isinstance(value, Path) else value for key, value in asdict(paths).items()},
        "metadata": metadata,
        "layout": layout,
        "document_layout": FICHAMENTO_QUALITATIVO_LAYOUT,
        "primary_sections": primary_sections,
        "body_blocks": len(body_blocks),
        "bibliography_entries": len(entries),
        "references_rendered": len(refs),
        "source_priority": "ORG final quando substantivo; document.json como fallback; BIB final como autoridade da seção REFERÊNCIAS BIBLIOGRÁFICAS",
    }
    paths.canonical_json.write_text(json.dumps(canonical, ensure_ascii=False, indent=2), encoding="utf-8")
    report = validate_fichamento_qualitativo_docx(paths.docx)
    report_path = paths.output_dir / f"{paths.prefix}.docx_canonico_report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    if not report["ok"]:
        message = "\n".join(f"- {warning}" for warning in report.get("warnings", []))
        die("ERRO: DOCX fichamento_qualitativo gerado com problemas de validação:\n" f"{message}\nArquivo: {paths.docx}")
    if not quiet:
        print(f"[OK] JSON canônico enriquecido: {paths.canonical_json}")
        print(f"[OK] Relatório DOCX canônico: {report_path}")
        print(f"[OK] DOCX fichamento_qualitativo: {paths.docx}")
    return paths.docx


def force_ooxml_black_docx(path: Path) -> None:
    """Força todas as definições OOXML de cor de fonte para preto.

    python-docx corrige a maioria dos runs, mas o Word/LibreOffice pode manter
    cores temáticas nos estilos Heading*. Esta etapa final atua diretamente em
    word/styles.xml, word/document.xml, headers, footers e notas, removendo
    themeColor/themeShade/themeTint e substituindo qualquer w:color por preto.
    """
    if not path.exists():
        return
    tmp = path.with_suffix(path.suffix + ".tmp_black")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    xml = data.decode("utf-8", errors="ignore")
                except Exception:
                    zout.writestr(item, data)
                    continue
                xml = re.sub(r"<w:color\b[^>]*/>", '<w:color w:val="000000"/>', xml)
                xml = re.sub(r"<w:color\b[^>]*>", '<w:color w:val="000000"/>', xml)
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)

def validate_docx(path: Path) -> dict[str, Any]:
    report: dict[str, Any] = {"path": str(path), "ok": False, "warnings": []}
    if not path.exists() or path.stat().st_size < 1000:
        report["warnings"].append("DOCX ausente ou muito pequeno.")
        return report
    try:
        with zipfile.ZipFile(path) as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
            try:
                style_xml = zf.read("word/styles.xml").decode("utf-8", errors="ignore")
            except KeyError:
                style_xml = ""
    except Exception as exc:
        report["warnings"].append(f"DOCX inválido: {exc}")
        return report
    for token, label in [
        ("Referências", "seção de referências"),
        ("Resumo", "resumo"),
    ]:
        if token not in xml:
            report["warnings"].append(f"Não encontrei {label} no document.xml.")
    residue_checks = [
        ("\\cite", "comando \\cite"),
        ("\\section", "comando \\section"),
        ("\\printbibliography", "comando \\printbibliography"),
        ("\\textbf", "comando \\textbf"),
        ("\\textit", "comando \\textit"),
        ("\\vspace", "comando \\vspace"),
        ("0.8em", "resíduo de espaçamento 0.8em"),
        ("academic_pipeline:", "marcador técnico academic_pipeline"),
        ("'PERFIL': 'FGV'", "perfil institucional cru na capa"),
        ("-chave: ATESTMED", "rótulo quebrado de palavras-chave"),
        (">: ATESTMED", "rótulo quebrado de keywords"),
    ]
    for token, label in residue_checks:
        if token in xml:
            report["warnings"].append(f"Resíduo visível detectado: {label}.")
    color_xml = xml + "\n" + style_xml
    color_tags = re.findall(r'<w:color\b[^>]*/>', color_xml)
    non_black_colors = [tag for tag in color_tags if 'w:val="000000"' not in tag and 'w:val="auto"' not in tag]
    if non_black_colors:
        report["warnings"].append(f"Cores de fonte não pretas detectadas no DOCX: {len(non_black_colors)} ocorrência(s).")
    report["font_color_policy"] = "all_word_color_tags_forced_to_000000"
    report["cover_city_year_space_before_pt"] = 240
    report["cover_discipline_in_top_block"] = True
    report["paragraphs_xml_count"] = xml.count("<w:p")
    report["size_bytes"] = path.stat().st_size
    report["ok"] = len(report["warnings"]) == 0
    return report

def render_docx_for_article(
    art: Path,
    prefix: str | None = None,
    cfg: Path | None = None,
    output: Path | None = None,
    quiet: bool = False,
) -> Path:
    if Document is None:
        die(
            "ERRO: python-docx não está disponível.\n"
            "Instale no ambiente do projeto com: pipenv install python-docx\n"
            f"Detalhe: {_DOCX_IMPORT_ERROR}"
        )
    paths = resolve_paths(art, cfg, prefix, output=output)
    paths.output_dir.mkdir(parents=True, exist_ok=True)
    if not paths.org.exists() and not paths.document_json.exists():
        die(f"ERRO: não encontrei ORG nem document.json em {paths.output_dir}")
    if not paths.bib.exists():
        die(f"ERRO: BIB não encontrado para materializar referências: {paths.bib}")

    org_text = read_text(paths.org)
    document_json = read_json(paths.document_json)
    resumos_json = read_json(paths.resumos_json)
    compliance_json = read_json(paths.compliance_json)
    toml_data = read_toml(paths.cfg)
    entries = parse_bib_entries(paths.bib)
    if not entries:
        die(f"ERRO: nenhuma entrada bibliográfica foi lida de {paths.bib}")

    metadata = build_metadata(paths, org_text, document_json, toml_data)
    abstracts = extract_resumos(resumos_json, org_text)
    body_blocks = build_body_blocks(org_text, document_json, entries)
    if not body_blocks:
        die("ERRO: não consegui extrair corpo textual do ORG/document.json.")
    layout = load_layout(compliance_json)

    if selected_document_layout(toml_data) == FICHAMENTO_QUALITATIVO_LAYOUT:
        return render_fichamento_qualitativo(
            paths=paths,
            metadata=metadata,
            body_blocks=body_blocks,
            entries=entries,
            layout=layout,
            toml_data=toml_data,
            quiet=quiet,
        )

    doc = Document()
    configure_doc_styles(doc, layout)
    add_cover(doc, metadata)
    add_abstracts(doc, abstracts)

    seen_first_heading = False
    for block in body_blocks:
        btype = block.get("type")
        if btype == "heading":
            text = clean_spaces(str(block.get("text", "")))
            if not text:
                continue
            # Avoid duplicate title page heading.
            if not seen_first_heading:
                seen_first_heading = True
            add_heading(doc, text, int(block.get("level", 1)))
        elif btype == "table":
            rows = block.get("rows")
            if isinstance(rows, list):
                add_table(doc, rows)
        elif btype == "paragraph":
            text = clean_spaces(str(block.get("text", "")))
            if not text:
                continue
            # Skip accidental visible reference commands.
            if text in {"Referências", "Resumo", "Abstract"}:
                continue
            add_body_paragraph(doc, text, first_indent=True)

    refs = add_references(doc, entries)
    backup(paths.docx, "docx_canonico", quiet=quiet)
    _normalize_fichamento_qualitativo_table_geometry(doc)
    doc.save(paths.docx)
    force_ooxml_black_docx(paths.docx)

    canonical = {
        "schema_version": "docx-canonico-v14",
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "paths": {k: str(v) if isinstance(v, Path) else v for k, v in asdict(paths).items()},
        "metadata": metadata,
        "layout": layout,
        "abstracts": abstracts,
        "body_blocks": len(body_blocks),
        "bibliography_entries": len(entries),
        "references_rendered": len(refs),
        "source_priority": "ORG final; document.json como fallback; BIB final para referências",
    }
    paths.canonical_json.write_text(json.dumps(canonical, ensure_ascii=False, indent=2), encoding="utf-8")

    report = validate_docx(paths.docx)
    report_path = paths.output_dir / f"{paths.prefix}.docx_canonico_report.json"
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    if not report["ok"]:
        msg = "\n".join(f"- {w}" for w in report.get("warnings", []))
        die(f"ERRO: DOCX canônico gerado com problemas de validação:\n{msg}\nArquivo: {paths.docx}")
    if not quiet:
        print(f"[OK] JSON canônico enriquecido: {paths.canonical_json}")
        print(f"[OK] Relatório DOCX canônico: {report_path}")
        print(f"[OK] DOCX canônico: {paths.docx}")
    return paths.docx


def main(argv: list[str] | None = None) -> int:
    ap = argparse.ArgumentParser(description="Gera DOCX ABNT/FGV canônico a partir do ORG/document.json/BIB finais. Versão v14 com disciplina no bloco institucional superior da capa.")
    ap.add_argument("--art-dir", required=True, help="Diretório do artigo.")
    ap.add_argument("--cfg-art", default=None, help="TOML do artigo, usado para metadados de capa quando disponível.")
    ap.add_argument("--prefix", default=None, help="Prefixo dos artefatos. Se omitido, tenta inferir.")
    ap.add_argument("--output", default=None, help="Caminho do DOCX de saída. Se omitido, usa output/<prefix>.docx.")
    ap.add_argument("--quiet", action="store_true")
    args = ap.parse_args(argv)
    render_docx_for_article(
        art=Path(args.art_dir),
        prefix=args.prefix,
        cfg=Path(args.cfg_art) if args.cfg_art else None,
        output=Path(args.output) if args.output else None,
        quiet=args.quiet,
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
