from __future__ import annotations

from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parents[2]
PIPELINE = ROOT / "app_bundle/scripts/pipeline"
if str(PIPELINE) not in sys.path:
    sys.path.insert(0, str(PIPELINE))

import render_docx_canonico as renderer

SECTIONS = list(renderer.FICHAMENTO_QUALITATIVO_SECTIONS)


def _entries():
    return {
        "a": {"author": "Braun, Virginia and Clarke, Victoria", "title": "Using thematic analysis in psychology", "journal": "Qualitative Research in Psychology", "year": "2006"},
        "b": {"author": "Souza, Luciana Karine de", "title": "Pesquisa com análise qualitativa de dados", "journal": "Arquivos Brasileiros de Psicologia", "year": "2019"},
    }


def _blocks():
    blocks = []
    for index, title in enumerate(SECTIONS):
        blocks.append({"type": "heading", "level": 1, "text": title})
        if index:
            blocks.append({"type": "paragraph", "text": "Parágrafo representativo do contrato visual do fichamento, com conteúdo suficiente para verificar estilo, alinhamento e ordenação sem depender de rede ou geração por IA."})
    return blocks


def test_selected_document_layout_prefers_documento_section():
    assert renderer.selected_document_layout({"documento": {"layout": "fichamento_qualitativo"}}) == "fichamento_qualitativo"


def test_fichamento_layout_machine_contract(tmp_path: Path):
    doc = renderer.Document()
    renderer.configure_doc_styles(doc, dict(renderer.DEFAULT_LAYOUT))
    renderer.add_fichamento_qualitativo_header(doc)
    renderer.add_fichamento_qualitativo_front_matter(
        doc,
        {"title": "Análise Temática em Pesquisa Qualitativa", "discipline": "Métodos Qualitativos", "professor": "Professor Teste", "author": "Pessoa de Teste", "year": "2026"},
        {"documento": {"data": "23 de agosto de 2026"}},
    )
    refs, seen = renderer.render_fichamento_qualitativo_body(doc, _blocks(), _entries())
    assert seen == SECTIONS
    assert len(refs) == 2
    target = tmp_path / "fichamento.docx"
    doc.save(target)
    report = renderer.validate_fichamento_qualitativo_docx(target)
    assert report["ok"], report["warnings"]
    assert report["header_table_count"] == 1
    assert report["header_shape"] == [1, 2]
    assert report["gradient_bands_below_header"] == 0
    assert report["compact_header"] is True
    assert report["technical_sheet_table_count"] == 1
    assert report["technical_sheet_labels"] == ["Disciplina", "Professor", "Aluno(a)", "Data"]
    assert abs(report["technical_sheet_label_column_width_cm"] - 3.1) <= 0.03
    assert abs(report["technical_sheet_value_column_width_cm"] - 12.9) <= 0.03
    assert report["page_break_count"] >= 7


def test_specialized_renderer_does_not_change_generic_dispatch_contract():
    source = (PIPELINE / "render_docx_canonico.py").read_text(encoding="utf-8")
    assert "force_ooxml_black_docx(paths.docx)" in source
    assert "if selected_document_layout(toml_data) == FICHAMENTO_QUALITATIVO_LAYOUT:" in source
    assert source.count("return render_fichamento_qualitativo(") == 1

# ---------------------------------------------------------------------------
# Cross-renderer final OOXML contract — user-approved golden geometry.
# ---------------------------------------------------------------------------
def test_fichamento_cross_renderer_raw_ooxml_contract(tmp_path):
    from docx import Document
    from app_bundle.scripts.pipeline import render_docx_canonico as renderer

    doc = Document()
    header = doc.sections[0].header.add_table(rows=1, cols=2, width=renderer.Cm(16))
    header.cell(0, 0).text = "FGV"
    header.cell(0, 1).text = "FICHAMENTO DE LEITURA"
    technical = doc.add_table(rows=4, cols=2)
    values = (
        ("Disciplina", "Métodos Qualitativos de Pesquisa em Ciência Política"),
        ("Professor", "Julio Cesar de Aguiar"),
        ("Aluno(a)", "Gustavo Magalhães Mendes de Tarso"),
        ("Data", "24 de agosto de 2026"),
    )
    for row, pair in zip(technical.rows, values):
        row.cells[0].text, row.cells[1].text = pair

    applied = renderer._normalize_fichamento_qualitativo_table_geometry(doc)
    assert applied == {"header": True, "technical": True}
    path = tmp_path / "cross_renderer_contract.docx"
    doc.save(path)
    report = renderer._validate_fichamento_cross_renderer_ooxml(path)
    assert report["ok"] is True
    assert report["header"]["layout"] == "fixed"
    assert report["header"]["tblW"] == 9071
    assert report["header"]["grid"] == [1588, 7483]
    assert [cell["tcW"] for cell in report["header"]["rows"][0]] == [1588, 7483]
    assert report["technical"]["layout"] == "fixed"
    assert report["technical"]["tblW"] == 9071
    assert report["technical"]["grid"] == [1758, 7313]
    for row in report["technical"]["rows"][:4]:
        assert [cell["tcW"] for cell in row] == [1758, 7313]
        assert row[0]["margins"] == {"top": 20, "bottom": 20, "start": 30, "end": 30}
        assert row[1]["margins"] == {"top": 20, "bottom": 20, "start": 60, "end": 60}
        assert row[0]["noWrap"] is True


def test_cross_renderer_normalizer_does_not_touch_unrelated_table():
    from docx import Document
    from app_bundle.scripts.pipeline import render_docx_canonico as renderer

    doc = Document()
    unrelated = doc.add_table(rows=1, cols=2)
    unrelated.cell(0, 0).text = "A"
    unrelated.cell(0, 1).text = "B"
    before = unrelated._tbl.xml
    applied = renderer._normalize_fichamento_qualitativo_table_geometry(doc)
    assert applied == {"header": False, "technical": False}
    assert unrelated._tbl.xml == before


def test_renderer_materializes_cross_renderer_geometry_before_save():
    import ast
    import inspect
    from app_bundle.scripts.pipeline import render_docx_canonico as renderer

    source = inspect.getsource(renderer.render_docx_for_article)
    tree = ast.parse(source)
    calls = []
    for node in ast.walk(tree):
        if isinstance(node, ast.Call):
            if isinstance(node.func, ast.Name):
                calls.append((node.lineno, node.func.id))
            elif isinstance(node.func, ast.Attribute):
                calls.append((node.lineno, node.func.attr))
    normalize_lines = [line for line, name in calls if name == "_normalize_fichamento_qualitativo_table_geometry"]
    save_lines = [line for line, name in calls if name == "save"]
    assert len(normalize_lines) == 1
    assert save_lines
    assert normalize_lines[0] < min(save_lines)



def test_fichamento_live_builders_materialize_exact_dxa_geometry():
    import ast
    import inspect
    from app_bundle.scripts.pipeline import render_docx_canonico as renderer

    expected = (
        (renderer.add_fichamento_qualitativo_header, "_materialize_fichamento_header_geometry"),
        (renderer.add_fichamento_qualitativo_front_matter, "_materialize_fichamento_technical_geometry"),
    )
    for function, helper in expected:
        tree = ast.parse(inspect.getsource(function))
        calls = [
            node for node in ast.walk(tree)
            if isinstance(node, ast.Call)
            and isinstance(node.func, ast.Name)
            and node.func.id == helper
        ]
        assert len(calls) == 1
