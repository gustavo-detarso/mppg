from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from institution_layouts import available_layouts, resolve_layout_spec
from paper_abstracts import (
    _bundle_items,
    _collect_source_strings,
    _labels_for_language,
    _source_for_abstract,
    abstract_sidecar_path,
    inject_paper_abstracts_into_docx,
    inject_paper_abstracts_into_org,
    main_document_abstract_languages,
    paper_abstracts_enabled,
    read_paper_abstract_bundle,
    requested_abstract_languages,
    write_paper_abstract_bundle,
)


class _Dumpable:
    def __init__(self, payload):
        self.payload = payload

    def model_dump(self, mode="python"):
        return self.payload


def _bundle() -> dict:
    return {
        "items": {
            "pt-br": {
                "language_code": "pt-br",
                "language_label": "Português (Brasil)",
                "heading": "Resumo",
                "keywords_heading": "Palavras-chave",
                "abstract": "Resumo principal.",
                "keywords": ["política pública", "evidências"],
                "include_keywords": True,
            },
            "en": {
                "language_code": "en",
                "language_label": "English",
                "heading": "Abstract",
                "keywords_heading": "Keywords",
                "abstract": "English abstract.",
                "keywords": ["public policy", "evidence"],
                "include_keywords": False,
            },
        }
    }


def test_paper_abstracts_enabled_requires_exact_profile_and_flags() -> None:
    cfg = {
        "projeto": {"preset": "paper_local_fgv"},
        "resumos_paper": {"ativo": True, "gerar_resumo_principal": True},
    }

    assert paper_abstracts_enabled(cfg)
    assert not paper_abstracts_enabled({"projeto": {"preset": "paper_local_fgv"}})
    assert not paper_abstracts_enabled(
        {
            "projeto": {"preset": "outro"},
            "resumos_paper": {"ativo": True, "gerar_resumo_principal": True},
        }
    )


def test_requested_abstract_languages_deduplicates_and_preserves_order() -> None:
    cfg = {
        "resumos_paper": {
            "principal": "pt-BR",
            "gerar_resumo_adicional": True,
            "idiomas_adicionais": ["en", "pt-BR", "es"],
        }
    }

    assert requested_abstract_languages(cfg) == [
        ("pt-br", "pt-BR"),
        ("en", "English"),
        ("es", "Español"),
    ]
    assert main_document_abstract_languages(cfg) == ["pt-br", "en", "es"]


def test_language_labels_use_known_and_generic_fallbacks() -> None:
    assert _labels_for_language("en", "English") == ("Abstract", "Keywords")
    assert _labels_for_language("pt-br", "Português") == ("Resumo", "Palavras-chave")
    assert _labels_for_language("nl", "Nederlands") == (
        "Resumo (Nederlands)",
        "Palavras-chave",
    )


def test_collect_source_strings_excludes_bibliography_and_diagnostics() -> None:
    payload = {
        "metadata": {"titulo": "Título acadêmico"},
        "sections": [{"title": "Introdução", "text": "Texto substantivo da análise."}],
        "bibliography": {"entries_used": ["silva2020"], "title": "Não coletar"},
        "diagnostics": {"warnings": ["pipeline interno"]},
        "figure": {"path": "/tmp/imagem.png", "title": "Não coletar por path protegido"},
    }

    values = _collect_source_strings(payload)

    assert "Título acadêmico" in values
    assert "Introdução" in values
    assert "Texto substantivo da análise." in values
    assert "Não coletar" not in values
    assert "pipeline interno" not in values
    assert "/tmp/imagem.png" not in values


def test_source_for_abstract_truncates_with_middle_and_final_markers() -> None:
    payload = {
        "sections": [
            {"title": "Introdução", "text": ("Texto inicial substantivo. " * 30)},
            {"title": "Meio", "text": ("Texto intermediário substantivo. " * 30)},
            {"title": "Conclusão", "text": ("Texto final substantivo. " * 30)},
        ]
    }

    source = _source_for_abstract(_Dumpable(payload), max_chars=300)

    assert "[trecho intermediário do paper]" in source
    assert "[trecho final do paper]" in source
    assert source.startswith("Introdução")
    assert source.endswith("Texto final substantivo.")


def test_bundle_items_filters_unknown_and_empty_rows() -> None:
    bundle = _bundle()
    bundle["items"]["es"] = {"abstract": ""}

    rows = _bundle_items(bundle, ["en", "es", "pt-BR", "fr"])

    assert [row["language_code"] for row in rows] == ["en", "pt-br"]


def test_sidecar_round_trip_is_utf8_and_validated(tmp_path: Path) -> None:
    path = abstract_sidecar_path(tmp_path, "artigo")

    write_paper_abstract_bundle(path, _bundle())
    restored = read_paper_abstract_bundle(path)

    assert path.name == "artigo.resumos_paper.json"
    assert restored == _bundle()
    assert "política pública" in path.read_text(encoding="utf-8")


def test_org_injection_is_idempotent_and_precedes_first_heading(tmp_path: Path) -> None:
    path = tmp_path / "artigo.org"
    path.write_text("#+TITLE: Teste\n* Introdução\nTexto.\n", encoding="utf-8")

    first = inject_paper_abstracts_into_org(path, _bundle(), ["pt-BR", "en"])
    second = inject_paper_abstracts_into_org(path, _bundle(), ["pt-BR", "en"])

    assert first.count("academic_pipeline:paper_abstracts:start") == 1
    assert second.count("academic_pipeline:paper_abstracts:start") == 1
    assert second.index("* Resumo") < second.index("* Introdução")
    assert "*Abstract:*" not in second
    assert "* Abstract" in second
    assert "Keywords:" not in second  # include_keywords=False no inglês


def test_org_injection_respects_native_front_matter_marker(tmp_path: Path) -> None:
    path = tmp_path / "artigo.org"
    original = (
        "#+BEGIN_COMMENT\n"
        "academic_pipeline:paper_abstracts:native\n"
        "#+END_COMMENT\n"
        "* Introdução\nTexto.\n"
    )
    path.write_text(original, encoding="utf-8")

    result = inject_paper_abstracts_into_org(path, _bundle(), ["pt-BR"])

    assert result == original
    assert path.read_text(encoding="utf-8") == original


def test_docx_injection_places_abstract_before_introduction(tmp_path: Path) -> None:
    path = tmp_path / "artigo.docx"
    document = Document()
    document.add_paragraph("Capa", style="Title")
    document.add_heading("Introdução", level=1)
    document.add_paragraph("Texto do corpo.")
    document.save(path)

    inject_paper_abstracts_into_docx(path, _bundle(), ["pt-BR"])

    restored = Document(path)
    texts = [p.text for p in restored.paragraphs]
    assert texts.index("Resumo") < texts.index("Introdução")
    assert texts.index("Resumo principal.") < texts.index("Introdução")
    assert "Palavras-chave: política pública; evidências." in texts


def test_available_layouts_filters_non_mapping_entries() -> None:
    cfg = {
        "__institution_profile__": {
            "layouts": {
                "paper_fgv": {"genero_academico": "paper"},
                "invalido": "texto",
            }
        }
    }

    assert available_layouts(cfg) == {
        "paper_fgv": {"genero_academico": "paper"}
    }


def test_layout_resolution_prefers_content_type_default() -> None:
    cfg = {
        "__institution_profile_name__": "fgv",
        "__institution_profile__": {
            "document_content_types": {
                "resumo_artigos": {"default_layout": "atividade_fgv"}
            },
            "document_types": {
                "atividade": {"default_layout": "outro"}
            },
            "layouts": {
                "atividade_fgv": {
                    "genero_academico": "atividade",
                    "front_matter": "atividade_fgv",
                    "classe_latex": "fgv-paper",
                    "template": "profile://atividade.org",
                    "validator": "fgv",
                }
            },
        },
        "documento": {
            "tipo_conteudo": "article_summary",
            "genero_academico": "atividade acadêmica",
        },
    }

    spec = resolve_layout_spec(cfg)

    assert spec.id == "atividade_fgv"
    assert spec.institution == "fgv"
    assert spec.genero_academico == "atividade"
    assert spec.tipo_conteudo == "resumo_artigos"
    assert spec.front_matter == "atividade_fgv"
    assert spec.classe_latex == "fgv-paper"
    assert spec.template == "profile://atividade.org"
    assert spec.validator == "fgv"


def test_layout_resolution_builds_legacy_fallback_for_undeclared_layout() -> None:
    cfg = {
        "instituicao": {"perfil": "fgv"},
        "documento": {
            "tipo_documento": "dissertação",
            "layout": "custom",
            "template_org": "modelo.org",
        },
    }

    spec = resolve_layout_spec(cfg)

    assert spec.id == "custom"
    assert spec.institution == "fgv"
    assert spec.genero_academico == "dissertacao"
    assert spec.classe_latex == "fgv-dissertacao"
    assert spec.template == "modelo.org"
