#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Resumo acadêmico multilíngue para o perfil ``paper_local_fgv``.

A camada opera *depois* da geração e validação do ``document.json`` canônico.
Ela produz um resumo no idioma principal do paper e, opcionalmente, versões
adicionais do resumo. Não altera o corpus, a bibliografia nem a análise do
conteúdo. O resultado é registrado em JSON e inserido nos artefatos ORG, PDF e
DOCX já produzidos pelo pipeline.
"""
from __future__ import annotations

import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

# Compatibilidade temporária entre pacote oficial e execução direta.
if __package__:
    from .document_translation import (
        normalize_language,
        requested_translation_languages,
    )
else:
    from document_translation import normalize_language, requested_translation_languages


class PaperAbstractError(RuntimeError):
    """Indica falha na geração, leitura ou inserção de resumos do paper."""


_LANGUAGE_LABELS: dict[str, tuple[str, str]] = {
    "pt": ("Resumo", "Palavras-chave"),
    "pt-br": ("Resumo", "Palavras-chave"),
    "en": ("Abstract", "Keywords"),
    "es": ("Resumen", "Palabras clave"),
    "fr": ("Résumé", "Mots-clés"),
    "it": ("Riassunto", "Parole chiave"),
    "de": ("Zusammenfassung", "Schlagwörter"),
}

_PROTECTED_PATH_TOKENS = {
    "bibliography", "bibliografia", "diagnostics", "diagnosticos",
    "citation", "citations", "citacao", "citacoes", "references",
    "referencias", "entries_used", "bib_path", "doi", "url", "uri",
    "href", "link", "path", "paths", "arquivo", "filename", "file",
    "image", "imagem", "figure_path", "source_info", "mindmap",
    "latex", "org", "json", "id", "identifier", "key", "keys",
    "slug", "created_at", "updated_at", "timestamp", "model", "version",
}


def _section(cfg: dict[str, Any], name: str) -> dict[str, Any]:
    value = cfg.get(name, {})
    return value if isinstance(value, dict) else {}


def _as_list(value: Any) -> list[str]:
    if isinstance(value, str):
        return [value]
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def paper_abstracts_enabled(cfg: dict[str, Any]) -> bool:
    """Retorna se o TOML solicita resumos do paper local."""
    project = _section(cfg, "projeto")
    section = _section(cfg, "resumos_paper")
    return (
        str(project.get("preset") or "").strip() == "paper_local_fgv"
        and bool(section.get("ativo", False))
        and bool(section.get("gerar_resumo_principal", True))
    )


def abstract_sidecar_path(out_dir: Path, prefix: str) -> Path:
    return out_dir / f"{prefix}.resumos_paper.json"


def requested_abstract_languages(
    cfg: dict[str, Any],
    *,
    include_full_translation_languages: bool = False,
) -> list[tuple[str, str]]:
    """Resolve os idiomas exigidos para a geração do resumo.

    O paper principal recebe apenas o idioma principal e os idiomas adicionais
    solicitados expressamente. Quando há uma versão integral traduzida, o seu
    idioma também é gerado para que essa versão tenha resumo coerente, sem
    transformar automaticamente o paper principal em bilíngue.
    """
    section = _section(cfg, "resumos_paper")
    principal_raw = section.get("principal") or _section(cfg, "idiomas_saida").get("principal") or "pt-BR"
    principal_code, principal_label = normalize_language(principal_raw)
    result: list[tuple[str, str]] = [(principal_code, principal_label)]
    seen: set[str] = {principal_code}

    if bool(section.get("gerar_resumo_adicional", False)):
        for raw in _as_list(section.get("idiomas_adicionais")):
            code, label = normalize_language(raw)
            if code not in seen:
                result.append((code, label))
                seen.add(code)

    if include_full_translation_languages:
        for code, label in requested_translation_languages(cfg):
            if code not in seen:
                result.append((code, label))
                seen.add(code)
    return result


def main_document_abstract_languages(cfg: dict[str, Any]) -> list[str]:
    """Idiomas que devem aparecer no paper principal, não nas cópias traduzidas."""
    return [code for code, _label in requested_abstract_languages(cfg, include_full_translation_languages=False)]


def _labels_for_language(code: str, label: str) -> tuple[str, str]:
    key = str(code or "").casefold()
    if key in _LANGUAGE_LABELS:
        return _LANGUAGE_LABELS[key]
    primary = key.split("-", 1)[0]
    if primary in _LANGUAGE_LABELS:
        return _LANGUAGE_LABELS[primary]
    return (f"Resumo ({label})", "Palavras-chave")


def _is_protected(path: tuple[str | int, ...]) -> bool:
    tokens = {str(item).casefold() for item in path if isinstance(item, str)}
    return bool(tokens & _PROTECTED_PATH_TOKENS)


def _looks_like_text(value: str) -> bool:
    stripped = value.strip()
    if len(stripped) < 3 or not re.search(r"[A-Za-zÀ-ÖØ-öø-ÿ]", stripped):
        return False
    if re.fullmatch(r"[A-Za-z0-9_.:/?=&%#@+\-]+", stripped):
        return False
    return True


def _collect_source_strings(payload: Any) -> list[str]:
    """Extrai redação substantiva, excluindo bibliografia e metadados técnicos."""
    values: list[str] = []

    def walk(value: Any, path: tuple[str | int, ...]) -> None:
        if _is_protected(path):
            return
        if isinstance(value, dict):
            for key, nested in value.items():
                walk(nested, (*path, str(key)))
            return
        if isinstance(value, list):
            for index, nested in enumerate(value):
                walk(nested, (*path, index))
            return
        if isinstance(value, str) and _looks_like_text(value):
            values.append(re.sub(r"\s+", " ", value).strip())

    walk(payload, ())
    # Remove duplicações frequentes entre metadados, títulos e a estrutura.
    seen: set[str] = set()
    return [item for item in values if not (item in seen or seen.add(item))]


def _source_for_abstract(document: Any, *, max_chars: int = 65000) -> str:
    try:
        payload = document.model_dump(mode="python")
    except TypeError:
        payload = document.model_dump()
    pieces = _collect_source_strings(payload)
    source = "\n\n".join(pieces).strip()
    if not source:
        raise PaperAbstractError("Não encontrei conteúdo textual suficiente no document.json para gerar o resumo.")
    if len(source) <= max_chars:
        return source
    # Preserva começo, miolo e encerramento para não reduzir a síntese apenas à introdução.
    first = max_chars * 45 // 100
    middle = max_chars * 20 // 100
    last = max_chars - first - middle
    midpoint = max(0, len(source) // 2 - middle // 2)
    return (
        source[:first]
        + "\n\n[trecho intermediário do paper]\n\n"
        + source[midpoint: midpoint + middle]
        + "\n\n[trecho final do paper]\n\n"
        + source[-last:]
    )


def _extract_response_content(response: Any) -> str:
    try:
        content = response.choices[0].message.content
    except Exception as exc:  # pragma: no cover - depende da resposta remota
        raise PaperAbstractError("A IA não devolveu conteúdo utilizável para o resumo.") from exc
    if isinstance(content, list):
        content = "".join(
            str(getattr(item, "text", "") or (item.get("text", "") if isinstance(item, dict) else ""))
            for item in content
        )
    return str(content or "").strip()


def _request_abstracts(
    client: Any,
    model: str,
    languages: list[tuple[str, str]],
    source: str,
    *,
    max_words: int,
) -> dict[str, dict[str, Any]]:
    requested = [{"code": code, "language": label} for code, label in languages]
    instructions = (
        "Você é um redator acadêmico. Com base exclusivamente no paper fornecido, produza um resumo "
        "acadêmico para cada idioma solicitado. Cada resumo deve sintetizar problema, objetivo, método ou "
        "base de evidências, argumento/resultado central e conclusão. Não cite autores, não invente dados, "
        "não use referências, DOI, URLs ou citações. Produza de 150 a " + str(max_words) + " palavras por idioma. "
        "Para cada idioma, forneça de 3 a 5 palavras-chave acadêmicas. Retorne JSON estrito no formato "
        '{"abstracts":[{"language_code":"...","abstract":"...","keywords":["..."]}]}. '
        "Cada language_code solicitado deve aparecer uma única vez."
    )
    payload = json.dumps({"languages": requested, "paper": source}, ensure_ascii=False)
    messages = [
        {"role": "system", "content": instructions},
        {"role": "user", "content": payload},
    ]
    try:
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            response_format={"type": "json_object"},
        )
    except Exception:
        response = client.chat.completions.create(model=model, messages=messages)
    content = _extract_response_content(response)
    try:
        parsed = json.loads(content)
    except json.JSONDecodeError as exc:
        raise PaperAbstractError("A IA devolveu o resumo fora do formato JSON esperado.") from exc
    rows = parsed.get("abstracts", []) if isinstance(parsed, dict) else []
    if not isinstance(rows, list):
        raise PaperAbstractError("A IA não devolveu a lista abstracts esperada.")

    by_code: dict[str, dict[str, Any]] = {}
    for row in rows:
        if not isinstance(row, dict):
            continue
        code, _label = normalize_language(row.get("language_code") or "")
        abstract = re.sub(r"\s+", " ", str(row.get("abstract") or "")).strip()
        keywords = [re.sub(r"\s+", " ", str(item)).strip() for item in _as_list(row.get("keywords"))]
        keywords = [item for item in keywords if item]
        if abstract and keywords:
            by_code[code] = {"abstract": abstract, "keywords": keywords[:8]}

    missing = [code for code, _label in languages if code not in by_code]
    if missing:
        raise PaperAbstractError("A IA não devolveu resumo para todos os idiomas: " + ", ".join(missing))
    return by_code


def generate_paper_abstract_bundle(client: Any, model: str, document: Any, cfg: dict[str, Any]) -> dict[str, Any]:
    """Gera e estrutura os resumos do paper principal e das versões integrais."""
    if not paper_abstracts_enabled(cfg):
        return {}
    section = _section(cfg, "resumos_paper")
    languages = requested_abstract_languages(cfg, include_full_translation_languages=True)
    try:
        max_words = int(section.get("max_palavras", 250))
    except (TypeError, ValueError):
        max_words = 250
    max_words = max(150, min(max_words, 350))
    source = _source_for_abstract(document)
    generated = _request_abstracts(client, model, languages, source, max_words=max_words)
    additional_keywords = bool(section.get("gerar_palavras_chave_adicionais", True))
    items: dict[str, dict[str, Any]] = {}
    main_code, _main_label = languages[0]
    for code, label in languages:
        heading, keywords_label = _labels_for_language(code, label)
        row = dict(generated[code])
        row.update(
            {
                "language_code": code,
                "language_label": label,
                "heading": heading,
                "keywords_heading": keywords_label,
                "include_keywords": True if code == main_code else additional_keywords,
            }
        )
        items[code] = row
    return {
        "schema_version": "1.0",
        "generated_at": datetime.now(timezone.utc).isoformat(timespec="seconds"),
        "model": model,
        "source_method": "síntese do document.json canônico; sem nova análise do corpus",
        "items": items,
    }


def read_paper_abstract_bundle(path: Path) -> dict[str, Any]:
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except Exception as exc:
        raise PaperAbstractError(f"Não foi possível ler o arquivo de resumos: {path}") from exc
    if not isinstance(payload, dict) or not isinstance(payload.get("items"), dict):
        raise PaperAbstractError("O arquivo de resumos possui formato inválido.")
    return payload


def write_paper_abstract_bundle(path: Path, bundle: dict[str, Any]) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(bundle, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return path


def _bundle_items(bundle: dict[str, Any], language_codes: Iterable[str]) -> list[dict[str, Any]]:
    items = bundle.get("items", {}) if isinstance(bundle, dict) else {}
    if not isinstance(items, dict):
        return []
    output: list[dict[str, Any]] = []
    for raw in language_codes:
        try:
            code, _label = normalize_language(raw)
        except PaperAbstractError:
            continue
        row = items.get(code)
        if isinstance(row, dict) and str(row.get("abstract") or "").strip():
            output.append(row)
    return output


_ORG_START = "#+BEGIN_COMMENT\nacademic_pipeline:paper_abstracts:start\n#+END_COMMENT"
_ORG_END = "#+BEGIN_COMMENT\nacademic_pipeline:paper_abstracts:end\n#+END_COMMENT"


def _org_block(rows: list[dict[str, Any]]) -> str:
    parts: list[str] = [_ORG_START, ""]
    for row in rows:
        heading = str(row.get("heading") or "Resumo").strip()
        abstract = str(row.get("abstract") or "").strip()
        parts.extend([
            f"* {heading}",
            ":PROPERTIES:",
            ":UNNUMBERED: t",
            ":END:",
            abstract,
        ])
        if bool(row.get("include_keywords", True)):
            keywords = "; ".join(str(item).strip() for item in _as_list(row.get("keywords")) if str(item).strip())
            if keywords:
                parts.extend(["", f"*{str(row.get('keywords_heading') or 'Palavras-chave').strip()}:* {keywords}."])
        parts.append("")
    parts.append(_ORG_END)
    return "\n".join(parts).rstrip() + "\n\n"


def inject_paper_abstracts_into_org(org_path: Path, bundle: dict[str, Any], language_codes: Iterable[str]) -> str:
    """Insere blocos de resumo antes da primeira seção do ORG, de modo idempotente."""
    rows = _bundle_items(bundle, language_codes)
    text = org_path.read_text(encoding="utf-8")
    # A v12 renderiza os resumos diretamente no front matter. Mantém a função
    # para retrocompatibilidade, mas não substitui o bloco nativo por headings.
    if "academic_pipeline:paper_abstracts:native" in text:
        return text
    if _ORG_START in text and _ORG_END in text:
        start = text.index(_ORG_START)
        end = text.index(_ORG_END, start) + len(_ORG_END)
        text = (text[:start] + text[end:]).lstrip("\n")
    if not rows:
        org_path.write_text(text.rstrip() + "\n", encoding="utf-8")
        return text
    block = _org_block(rows)
    match = re.search(r"(?m)^\*\s+", text)
    if match:
        text = text[:match.start()] + block + text[match.start():]
    else:
        text = text.rstrip() + "\n\n" + block
    org_path.write_text(text.rstrip() + "\n", encoding="utf-8")
    return text


def _docx_anchor(document: Any) -> Any | None:
    """Busca a primeira seção de conteúdo, preservando a capa do DOCX."""
    for paragraph in getattr(document, "paragraphs", []):
        text = str(getattr(paragraph, "text", "") or "").strip()
        style_name = str(getattr(getattr(paragraph, "style", None), "name", "") or "").casefold()
        if not text:
            continue
        if "heading" in style_name or "título" in style_name or "titulo" in style_name:
            if "title" not in style_name and "capa" not in style_name:
                return paragraph
        if re.match(r"^(?:\d+(?:\.\d+)*\s+)?(?:introdução|introduction|introduccion|introducción)\b", text, flags=re.I):
            return paragraph
    return None


def inject_paper_abstracts_into_docx(docx_path: Path, bundle: dict[str, Any], language_codes: Iterable[str]) -> None:
    """Insere os resumos antes da primeira seção do conteúdo no DOCX renderizado."""
    rows = _bundle_items(bundle, language_codes)
    if not rows:
        return
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - depende do ambiente do usuário
        raise PaperAbstractError("python-docx não está disponível para inserir resumos no DOCX.") from exc
    document = Document(str(docx_path))
    anchor = _docx_anchor(document)
    created: list[Any] = []
    for row in rows:
        heading = document.add_paragraph()
        heading_run = heading.add_run(str(row.get("heading") or "Resumo"))
        heading_run.bold = True
        created.append(heading)
        paragraph = document.add_paragraph(str(row.get("abstract") or "").strip())
        created.append(paragraph)
        if bool(row.get("include_keywords", True)):
            keywords = "; ".join(str(item).strip() for item in _as_list(row.get("keywords")) if str(item).strip())
            if keywords:
                keyword_paragraph = document.add_paragraph()
                keyword_paragraph.add_run(str(row.get("keywords_heading") or "Palavras-chave") + ": ").bold = True
                keyword_paragraph.add_run(keywords + ".")
                created.append(keyword_paragraph)
        created.append(document.add_paragraph())
    if anchor is not None:
        for paragraph in created:
            anchor._p.addprevious(paragraph._p)
    document.save(str(docx_path))
