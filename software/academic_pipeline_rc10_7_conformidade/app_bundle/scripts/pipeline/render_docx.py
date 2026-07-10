#!/usr/bin/env python3
# -*- coding: utf-8 -*-
from __future__ import annotations

import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any

from document_model import AcademicDocument, Citation, TextSpan, Block
from utils import normalize_title_loose, write_text
from render_org_latex import is_ai_generated_reference_section_title

try:
    from docx import Document
    from docx.shared import Cm, Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except Exception:  # pragma: no cover
    Document = None  # type: ignore


def parse_bib_entries(text: str) -> list[str]:
    entries = []
    i = 0
    while True:
        at = text.find("@", i)
        if at < 0:
            break
        brace = text.find("{", at)
        if brace < 0:
            break
        depth = 0
        for j in range(brace, len(text)):
            if text[j] == "{":
                depth += 1
            elif text[j] == "}":
                depth -= 1
                if depth == 0:
                    entries.append(text[at:j + 1].strip())
                    i = j + 1
                    break
        else:
            break
    return entries


def bib_key(entry: str) -> str | None:
    m = re.match(r"\s*@[^{]+\{\s*([^,]+)\s*,", entry, re.S)
    return m.group(1).strip() if m else None


def field(entry: str, name: str) -> str:
    m = re.search(rf"(?is)\b{re.escape(name)}\s*=\s*", entry)
    if not m:
        return ""
    i = m.end()
    while i < len(entry) and entry[i].isspace():
        i += 1
    if i >= len(entry) or entry[i] != "{":
        return ""
    depth, start = 0, i + 1
    for j in range(i, len(entry)):
        if entry[j] == "{":
            depth += 1
        elif entry[j] == "}":
            depth -= 1
            if depth == 0:
                return re.sub(r"\s+", " ", entry[start:j]).strip()
    return ""


def bib_to_reference_map(bib_path: Path | None) -> dict[str, dict[str, str]]:
    if not bib_path or not bib_path.exists():
        return {}
    out: dict[str, dict[str, str]] = {}
    for e in parse_bib_entries(bib_path.read_text(encoding="utf-8", errors="ignore")):
        k = bib_key(e)
        if not k:
            continue
        out[k] = {name: field(e, name) for name in ["author", "editor", "title", "year", "journaltitle", "booktitle", "publisher", "volume", "number", "pages", "doi", "url"]}
    return out


def first_author_label(author: str) -> str:
    if not author:
        return "Autor"
    first = re.split(r"\s+and\s+", author, flags=re.I)[0].strip()
    if "," in first:
        return first.split(",", 1)[0].strip()
    parts = first.split()
    return parts[-1] if parts else first


def citation_text(cit: Citation, refs: dict[str, dict[str, str]]) -> str:
    labels = []
    for key in cit.keys:
        meta = refs.get(key, {})
        author = first_author_label(meta.get("author") or meta.get("editor") or key)
        year = meta.get("year") or "s.d."
        labels.append(f"{author} ({year})" if cit.mode == "narrative" else f"{author}, {year}")
    if not labels:
        return ""
    if cit.mode == "narrative":
        if len(labels) == 1:
            return labels[0]
        return ", ".join(labels[:-1]) + " e " + labels[-1]
    return "(" + "; ".join(labels) + ")"


def reference_text(key: str, meta: dict[str, str]) -> str:
    author = (meta.get("author") or meta.get("editor") or "").replace(" and ", "; ")
    year = meta.get("year") or "s.d."
    title = meta.get("title") or key
    venue = meta.get("journaltitle") or meta.get("booktitle") or meta.get("publisher") or ""
    vol = meta.get("volume") or ""
    num = meta.get("number") or ""
    pages = meta.get("pages") or ""
    doi = meta.get("doi") or ""
    parts = []
    if author:
        parts.append(author + ".")
    parts.append(f"({year}).")
    parts.append(title + ".")
    if venue:
        parts.append(venue + (f", {vol}" if vol else "") + (f"({num})" if num else "") + (f", {pages}" if pages else "") + ".")
    if doi:
        parts.append("https://doi.org/" + doi)
    return " ".join(parts)


def setup_styles(doc: Any) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    for name, size, bold in [("Title", 16, True), ("Heading 1", 14, True), ("Heading 2", 13, True), ("Heading 3", 12, True)]:
        try:
            st = styles[name]
            st.font.name = "Times New Roman"
            st.font.size = Pt(size)
            st.font.bold = bold
        except Exception:
            pass


def _paper_abstract_sidecar_exists(output_path: Path, cfg: dict[str, Any] | None) -> bool:
    """Informa se o resumo auditável será inserido pelo pipeline após o DOCX.

    Quando existe sidecar, o ``document_model.abstract`` não deve ser emitido
    aqui: isso evitaria duplicação entre o resumo original do modelo e o pacote
    final multilíngue gerado a partir do document.json validado.
    """
    cfg = cfg or {}
    project = cfg.get("projeto", {}) if isinstance(cfg.get("projeto"), dict) else {}
    section = cfg.get("resumos_paper", {}) if isinstance(cfg.get("resumos_paper"), dict) else {}
    if str(project.get("preset") or "").strip() != "paper_local_fgv":
        return False
    if not bool(section.get("ativo", False)) or not bool(section.get("gerar_resumo_principal", True)):
        return False
    paths = cfg.get("paths", {}) if isinstance(cfg.get("paths"), dict) else {}
    prefix = str(paths.get("document_prefix") or "").strip() or output_path.stem
    for directory in [output_path.parent, *output_path.parents][:5]:
        if (directory / f"{prefix}.resumos_paper.json").is_file():
            return True
    return False


def add_paragraph_from_content(doc: Any, content: list[Any], refs: dict[str, dict[str, str]]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    p.alignment = 3  # justify
    for item in content:
        if isinstance(item, Citation) or getattr(item, "type", None) == "citation":
            p.add_run(citation_text(item, refs))
        else:
            r = p.add_run(getattr(item, "text", ""))
            r.italic = bool(getattr(item, "italic", False))
            r.bold = bool(getattr(item, "bold", False))


def _md_inlines(content: list[Any]) -> str:
    parts: list[str] = []
    for item in content:
        if isinstance(item, Citation) or getattr(item, "type", None) == "citation":
            keys = ["@" + k for k in getattr(item, "keys", [])]
            if getattr(item, "mode", "parenthetical") == "narrative" and keys:
                # Pandoc narrative citation is not as universal. Keep an author-in-text approximation.
                parts.append("[" + "; ".join(keys) + "]")
            else:
                parts.append("[" + "; ".join(keys) + "]")
        else:
            text = getattr(item, "text", "")
            if getattr(item, "bold", False):
                text = f"**{text}**"
            if getattr(item, "italic", False):
                text = f"*{text}*"
            parts.append(text)
    return "".join(parts)


def render_markdown_for_pandoc(doc_model: AcademicDocument, out_dir: Path, cfg: dict[str, Any] | None = None) -> str:
    meta = doc_model.metadata
    lines = [
        f"% {meta.titulo}",
        f"% {meta.autor}",
        f"% {meta.ano or meta.data}",
        "",
        meta.instituicao,
        "",
        meta.curso,
        "",
        meta.disciplina,
        "",
    ]
    if doc_model.abstract and doc_model.abstract.texto and not _paper_abstract_sidecar_exists(out_dir / "documento.docx", cfg):
        lines += ["# RESUMO", "", doc_model.abstract.texto, ""]
        if doc_model.abstract.palavras_chave:
            lines += ["Palavras-chave: " + "; ".join(doc_model.abstract.palavras_chave) + ".", ""]
    for sec in doc_model.sections:
        if is_ai_generated_reference_section_title(sec.title):
            continue
        lines += ["#" * max(1, min(sec.level, 6)) + " " + sec.title, ""]
        for block in sec.blocks:
            if block.type == "paragraph":
                lines += [_md_inlines(block.content), ""]
            elif block.type == "heading":
                lines += ["#" * max(1, min(block.level, 6)) + " " + (block.text or block.title), ""]
            elif block.type in {"bullet_list", "numbered_list"}:
                for i, item in enumerate(block.items, start=1):
                    lines.append((f"{i}. " if block.type == "numbered_list" else "- ") + item)
                lines.append("")
            elif block.type == "quote":
                lines += ["> " + block.text.replace("\n", "\n> "), ""]
            elif block.type == "figure" and block.path:
                lines += [f"![{block.title}]({block.path})", ""]
    lines += ["# REFERÊNCIAS", "", "::: {#refs}", ":::", ""]
    for fig in doc_model.figures:
        if fig.placement == "after_references":
            lines += ["\\newpage", "", "# " + fig.title, "", f"![{fig.title}]({fig.path})", ""]
    return "\n".join(lines)


def _render_docx_with_pandoc(doc_model: AcademicDocument, output_path: Path, bib_path: Path | None, reference_docx: Path | None, cfg: dict[str, Any] | None) -> Path | None:
    pandoc = shutil.which("pandoc")
    if not pandoc or not bib_path or not bib_path.exists():
        return None
    cfg = cfg or {}
    docx_cfg = cfg.get("docx", {}) if isinstance(cfg.get("docx"), dict) else {}
    config_dir = Path(str(cfg.get("__config_dir__") or output_path.parent)).resolve()
    raw_csl = docx_cfg.get("csl_path") or (cfg.get("bibliografia", {}) if isinstance(cfg.get("bibliografia"), dict) else {}).get("docx_csl")
    csl = Path(str(raw_csl)).expanduser()
    if raw_csl and not csl.is_absolute():
        csl = config_dir / csl
    with tempfile.TemporaryDirectory() as tmp:
        md_path = Path(tmp) / "documento.md"
        md_path.write_text(render_markdown_for_pandoc(doc_model, output_path.parent, cfg=cfg), encoding="utf-8")
        cmd = [pandoc, str(md_path), "-o", str(output_path), "--citeproc", "--bibliography", str(bib_path), "--resource-path", str(output_path.parent)]
        if csl.exists():
            cmd += ["--csl", str(csl)]
        if reference_docx and reference_docx.exists():
            cmd += ["--reference-doc", str(reference_docx)]
        proc = subprocess.run(cmd, cwd=str(output_path.parent), text=True, capture_output=True)
        if proc.returncode == 0 and output_path.exists():
            return output_path
        diag = output_path.with_suffix(".pandoc_docx_erro.txt")
        diag.write_text("CMD: " + " ".join(cmd) + "\n\nSTDOUT:\n" + proc.stdout + "\n\nSTDERR:\n" + proc.stderr, encoding="utf-8")
        return None


def render_docx(doc_model: AcademicDocument, output_path: Path, bib_path: Path | None = None, reference_docx: Path | None = None, cfg: dict[str, Any] | None = None) -> Path:
    cfg = cfg or {}
    docx_cfg = cfg.get("docx", {}) if isinstance(cfg.get("docx"), dict) else {}
    if bool(docx_cfg.get("usar_pandoc", False)):
        rendered = _render_docx_with_pandoc(doc_model, output_path, bib_path, reference_docx, cfg)
        if rendered:
            return rendered
        if bool(docx_cfg.get("falhar_se_pandoc_falhar", False)):
            raise RuntimeError(f"Falha ao renderizar DOCX via Pandoc. Veja {output_path.with_suffix('.pandoc_docx_erro.txt')}")
    if Document is None:
        raise RuntimeError("python-docx não está instalado.")
    doc = Document(str(reference_docx)) if reference_docx and reference_docx.exists() else Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Cm(3)
    sec.left_margin = Cm(3)
    sec.right_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    refs = bib_to_reference_map(bib_path)
    meta = doc_model.metadata

    if meta.tipo_documento == "atividade":
        doc.add_heading("FICHA TÉCNICA", level=1)
        table = doc.add_table(rows=8, cols=2)
        table.style = "Table Grid"
        rows = [("Curso", meta.curso), ("Turma", meta.turma), ("Pólo", meta.polo), ("Disciplina", meta.disciplina), ("Professor", meta.professor), ("Aluno(s)", meta.autor), ("Data", meta.data or meta.ano), ("Título do trabalho", meta.titulo)]
        for i, (a, b) in enumerate(rows):
            table.cell(i, 0).text = a
            table.cell(i, 1).text = b
        doc.add_page_break()
    else:
        for text in [meta.instituicao, meta.programa, meta.curso, meta.disciplina, meta.titulo, meta.autor, meta.tipo_trabalho, meta.professor, meta.cidade, meta.ano or meta.data]:
            if not text:
                continue
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            r.font.name = "Times New Roman"
            r.font.size = Pt(12)
            if text == meta.titulo:
                r.bold = True
                r.font.size = Pt(14)
        doc.add_page_break()

    if doc_model.abstract and doc_model.abstract.texto and not _paper_abstract_sidecar_exists(output_path, cfg):
        doc.add_heading(doc_model.abstract.titulo.upper(), level=1)
        p = doc.add_paragraph(doc_model.abstract.texto)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = 3
        if doc_model.abstract.palavras_chave:
            doc.add_paragraph("Palavras-chave: " + "; ".join(doc_model.abstract.palavras_chave) + ".")

    for sec_model in doc_model.sections:
        if is_ai_generated_reference_section_title(sec_model.title):
            continue
        doc.add_heading(sec_model.title.upper() if sec_model.level == 1 else sec_model.title, level=sec_model.level)
        for block in sec_model.blocks:
            if block.type == "paragraph":
                add_paragraph_from_content(doc, block.content, refs)
            elif block.type == "heading":
                doc.add_heading(block.text or block.title, level=block.level)
            elif block.type == "quote":
                p = doc.add_paragraph(block.text)
                p.paragraph_format.left_indent = Cm(4)
            elif block.type in {"bullet_list", "numbered_list"}:
                style = "List Bullet" if block.type == "bullet_list" else "List Number"
                for item in block.items:
                    doc.add_paragraph(item, style=style)
            elif block.type == "table" and block.table:
                rows = len(block.table.rows) + (1 if block.table.headers else 0)
                cols = len(block.table.headers or (block.table.rows[0] if block.table.rows else []))
                if rows and cols:
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = "Table Grid"
                    r = 0
                    if block.table.headers:
                        for c, h in enumerate(block.table.headers):
                            table.cell(0, c).text = h
                        r = 1
                    for row in block.table.rows:
                        for c, val in enumerate(row[:cols]):
                            table.cell(r, c).text = str(val)
                        r += 1
            elif block.type == "figure" and block.path:
                if block.page_break_before:
                    doc.add_page_break()
                img = Path(output_path.parent / block.path)
                if not img.exists():
                    img = Path(block.path)
                if img.exists():
                    doc.add_picture(str(img), width=Inches(6.2))
                if block.title:
                    cap = doc.add_paragraph(block.title)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif block.type == "page_break":
                doc.add_page_break()

    doc.add_heading("REFERÊNCIAS", level=1)
    used = doc_model.bibliography.entries_used or list(refs.keys())
    for key in used:
        if key in refs:
            p = doc.add_paragraph(reference_text(key, refs[key]))
            p.paragraph_format.left_indent = Cm(1.25)
            p.paragraph_format.first_line_indent = Cm(-1.25)

    for fig in doc_model.figures:
        if fig.placement == "after_references":
            if fig.page_break_before:
                doc.add_page_break()
            doc.add_heading(fig.title.upper(), level=1)
            img = Path(output_path.parent / fig.path)
            if not img.exists():
                img = Path(fig.path)
            if img.exists():
                doc.add_picture(str(img), width=Inches(6.4))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
