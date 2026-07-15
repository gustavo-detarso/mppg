from __future__ import annotations

import json
from pathlib import Path
from zipfile import ZipFile

import pytest

from docx import Document
from docx.shared import Cm, RGBColor

from render_docx_canonico import (
    ArticlePaths,
    DEFAULT_LAYOUT,
    add_cover,
    add_references,
    add_table,
    build_body_blocks,
    build_metadata,
    clean_latex_inline,
    clean_spaces,
    configure_doc_styles,
    extract_org_abstracts,
    extract_resumos,
    force_ooxml_black_docx,
    materialize_citations,
    parse_bib_entries,
    parse_document_json_blocks,
    parse_org_blocks,
    parse_org_latex_header_macros,
    parse_org_metadata,
    reference_line,
    render_docx_for_article,
    resolve_paths,
    validate_docx,
)


def _cm(value) -> float:
    return round(value.cm, 2)


def _paths(tmp_path: Path, prefix: str = "artigo") -> ArticlePaths:
    out = tmp_path / "output"
    return ArticlePaths(
        art_dir=tmp_path,
        output_dir=out,
        prefix=prefix,
        org=out / f"{prefix}.org",
        bib=out / f"{prefix}.bib",
        document_json=out / f"{prefix}.document.json",
        resumos_json=out / f"{prefix}.resumos_paper.json",
        compliance_json=out / f"{prefix}.compliance_report.json",
        run_report_json=out / f"{prefix}.run_report.json",
        cfg=tmp_path / f"{prefix}.toml",
        docx=out / f"{prefix}.docx",
        canonical_json=out / f"{prefix}.canonical_docx.json",
    )


def _bib_text() -> str:
    return (
        "@article{silva2020,\n"
        "  author={Silva, Ana and Souza, Bruno},\n"
        "  title={Políticas baseadas em evidências},\n"
        "  year={2020},\n"
        "  journaltitle={Revista de Políticas},\n"
        "  volume={10},\n"
        "  number={2},\n"
        "  pages={1--20},\n"
        "  doi={10.1000/teste}\n"
        "}\n"
    )


def test_clean_spaces_normalizes_unicode_and_line_breaks() -> None:
    assert clean_spaces(" A\xa0B \r\n C\u00ad \n\n\n D ") == "A B\nC\n\nD"


def test_clean_latex_inline_preserves_visible_formatting_content() -> None:
    raw = r"\textbf{Análise} com \emph{evidência}, \href{https://x}{fonte} e \LaTeX{}."

    cleaned = clean_latex_inline(raw)

    assert cleaned == "Análise com evidência, fonte (https://x) e LaTeX."


def test_resolve_paths_infers_prefix_from_document_json(tmp_path: Path) -> None:
    output = tmp_path / "output"
    output.mkdir()
    (output / "meu_artigo.document.json").write_text("{}", encoding="utf-8")

    paths = resolve_paths(tmp_path, None, None)

    assert paths.prefix == "meu_artigo"
    assert paths.org == output / "meu_artigo.org"
    assert paths.docx == output / "meu_artigo.docx"


def test_resolve_paths_honors_explicit_output_file(tmp_path: Path) -> None:
    target = tmp_path / "custom" / "resultado.docx"

    paths = resolve_paths(tmp_path, None, "prefixo", output=target)

    assert paths.output_dir == target.parent.resolve()
    assert paths.docx == target.resolve()
    assert paths.prefix == "prefixo"


def test_org_metadata_and_latex_macros_are_parsed() -> None:
    org = (
        "#+TITLE: Título final\n"
        "#+AUTHOR: Autora\n"
        "#+LATEX_HEADER: \\institution{FGV}\\coursename{Mestrado}\\professorname{Docente}\n"
    )

    assert parse_org_metadata(org) == {
        "title": "Título final",
        "author": "Autora",
        "latex_header": r"\institution{FGV}\coursename{Mestrado}\professorname{Docente}",
    }
    assert parse_org_latex_header_macros(org) == {
        "institution": "FGV",
        "coursename": "Mestrado",
        "professorname": "Docente",
    }


def test_metadata_prefers_final_org_and_deduplicates_program_course(tmp_path: Path) -> None:
    paths = _paths(tmp_path)
    org = (
        "#+TITLE: Título do ORG\n"
        "#+AUTHOR: Autora do ORG\n"
        "#+LATEX_HEADER: \\institution{Fundação Getúlio Vargas}\n"
        "#+LATEX_HEADER: \\programname{Mestrado em Políticas Públicas}\n"
        "#+LATEX_HEADER: \\coursename{Mestrado em Políticas Públicas}\n"
        "#+LATEX_HEADER: \\disciplinename{Decisões Baseadas em Evidências}\n"
        "#+LATEX_HEADER: \\professorname{Professor X}\n"
        "#+LATEX_HEADER: \\cityname{Brasília}\n"
    )

    metadata = build_metadata(
        paths,
        org,
        {"metadata": {"titulo": "JSON antigo", "ano": "2025"}},
        {"documento": {"titulo": "TOML antigo"}, "ano": "2026"},
    )

    assert metadata["title"] == "Título do ORG"
    assert metadata["author"] == "Autora do ORG"
    assert metadata["institution"] == "Fundação Getúlio Vargas"
    assert not (metadata["program"] and metadata["course"])
    assert metadata["discipline"] == "Decisões Baseadas em Evidências"
    assert metadata["professor"] == "Professor X"


def test_parse_bib_entries_materializes_year_from_date(tmp_path: Path) -> None:
    bib = tmp_path / "refs.bib"
    bib.write_text(
        "@book{x, author={Silva, Ana}, title={Livro}, date={2024-03-01}, publisher={Editora}}\n",
        encoding="utf-8",
    )

    entries = parse_bib_entries(bib)

    assert entries["x"]["entrytype"] == "book"
    assert entries["x"]["year"] == "2024"
    assert entries["x"]["publisher"] == "Editora"


def test_materialize_citations_handles_textual_parenthetical_and_org_syntax() -> None:
    entries = {
        "silva2020": {
            "id": "silva2020",
            "author": "Silva, Ana",
            "year": "2020",
            "title": "Título",
        }
    }
    text = (
        r"\textcite{silva2020} argumenta. "
        r"Outro achado \parencite{silva2020}. "
        r"Também [@silva2020]."
    )

    materialized = materialize_citations(text, entries)

    assert "Silva (2020) argumenta" in materialized
    assert materialized.count("(Silva, 2020)") == 2
    assert "\\textcite" not in materialized
    assert "\\parencite" not in materialized


def test_reference_line_follows_current_abnt_like_contract() -> None:
    rendered = reference_line(
        {
            "id": "silva2020",
            "author": "Silva, Ana and Souza, Bruno",
            "title": "Título",
            "year": "2020",
            "journaltitle": "Revista",
            "volume": "10",
            "number": "2",
            "pages": "1--20",
            "doi": "10.1000/teste",
        }
    )

    assert rendered.startswith("SILVA, Ana; SOUZA, Bruno. Título.")
    assert "Revista, v. 10, n. 2, p. 1--20, 2020." in rendered
    assert rendered.endswith("DOI: 10.1000/teste.")


def test_extract_org_abstracts_reads_export_latex_blocks() -> None:
    org = r"""
#+begin_export latex
\begingroup
\textbf{Resumo}
Resumo em português com conteúdo substantivo.
\noindent\textbf{Palavras-chave:} política pública; evidências; avaliação.
\vspace{0.8em}
\textbf{Abstract}
English abstract with substantive content.
\noindent\textbf{Keywords:} public policy; evidence; evaluation.
\endgroup
#+end_export
* Introdução
Texto.
"""

    result = extract_org_abstracts(org)

    assert result["abstract_pt"] == "Resumo em português com conteúdo substantivo."
    assert result["keywords_pt"] == ["política pública", "evidências", "avaliação"]
    assert result["abstract_en"] == "English abstract with substantive content."
    assert result["keywords_en"] == ["public policy", "evidence", "evaluation"]


@pytest.mark.xfail(
    strict=True,
    reason=(
        "Defeito legado catalogado: extract_org_abstracts inclui a linha "
        "Palavras-chave no abstract quando o ORG usa heading Resumo seguido "
        "de rótulo inline, sem heading próprio."
    ),
)
def test_extract_resumos_should_separate_inline_keywords_from_heading_abstract() -> None:
    sidecar = {
        "items": {
            "pt-br": {"abstract": "Resumo lateral.", "keywords": ["um", "dois"]},
            "en": {"abstract": "Sidecar abstract.", "keywords": ["one", "two"]},
        }
    }
    org = "* Resumo\nResumo final do ORG.\n\nPalavras-chave: org; final.\n"

    result = extract_resumos(sidecar, org)

    assert result["abstract_pt"] == "Resumo final do ORG."
    assert result["keywords_pt"] == ["org", "final"]
    assert result["abstract_en"] == "Sidecar abstract."


def test_parse_org_blocks_ignores_comments_exports_and_reference_tail() -> None:
    entries = {"silva2020": {"id": "silva2020", "author": "Silva, Ana", "year": "2020"}}
    org = (
        "#+TITLE: Teste\n"
        "#+BEGIN_COMMENT\nacademic_pipeline:interno\n#+END_COMMENT\n"
        "#+begin_export latex\n\\textbf{Resumo} escondido\n#+end_export\n"
        "* Introdução\n"
        "Texto com \\parencite{silva2020}.\n\n"
        "| Item | Valor |\n|------+-------|\n| A | 1 |\n"
        "* Referências\n"
        "Entrada que não deve virar corpo.\n"
    )

    blocks = parse_org_blocks(org, entries)

    assert blocks[0] == {"type": "heading", "level": 1, "text": "Introdução"}
    assert blocks[1]["type"] == "paragraph"
    assert "(Silva, 2020)" in blocks[1]["text"]
    assert blocks[2] == {"type": "table", "rows": [["Item", "Valor"], ["A", "1"]]}
    assert all("Entrada que não deve" not in str(block) for block in blocks)


def test_document_json_blocks_use_legacy_section_text_contract() -> None:
    document_json = {
        "sections": [
            {"title": "Introdução", "text": "\\section{Introdução}\n\nTexto um.\n\nTexto dois."}
        ]
    }

    blocks = parse_document_json_blocks(document_json, {})

    assert blocks == [
        {"type": "heading", "level": 1, "text": "Introdução"},
        {"type": "paragraph", "text": "Texto um."},
        {"type": "paragraph", "text": "Texto dois."},
    ]


def test_build_body_blocks_falls_back_to_document_json_for_short_org() -> None:
    org = "* Introdução\nTexto curto."
    document_json = {
        "sections": [
            {"title": "Introdução", "text": "Texto substancial do JSON."}
        ]
    }

    blocks = build_body_blocks(org, document_json, {})

    assert blocks[0]["text"] == "Introdução"
    assert blocks[1]["text"] == "Texto substancial do JSON."


def test_configure_doc_styles_sets_a4_margins_and_black_fonts() -> None:
    document = Document()

    configure_doc_styles(document, DEFAULT_LAYOUT)

    section = document.sections[0]
    assert _cm(section.page_width) == 21.0
    assert _cm(section.page_height) == 29.7
    assert _cm(section.top_margin) == 3.0
    assert _cm(section.left_margin) == 3.0
    assert _cm(section.right_margin) == 2.0
    assert _cm(section.bottom_margin) == 2.0
    assert document.styles["Normal"].font.name == "Times New Roman"
    assert str(document.styles["Heading 1"].font.color.rgb) == "000000"


def test_cover_deduplicates_repeated_institutional_lines() -> None:
    document = Document()
    metadata = {
        "institution": "Fundação Getúlio Vargas",
        "program": "Mestrado em Políticas Públicas",
        "course": "Mestrado em Políticas Públicas",
        "discipline": "Disciplina",
        "author": "Autora",
        "title": "Título",
        "subtitle": "",
        "covernote": "",
        "professor": "Professor X",
        "city": "Brasília",
        "year": "2026",
    }

    add_cover(document, metadata)

    texts = [p.text for p in document.paragraphs]
    assert texts.count("MESTRADO EM POLÍTICAS PÚBLICAS") == 1
    assert "FUNDAÇÃO GETÚLIO VARGAS" in texts
    assert "DISCIPLINA" in texts


def test_table_and_references_render_current_alignment_and_sorting() -> None:
    document = Document()

    add_table(document, [["Coluna A", "Coluna B"], ["A", "1"]])
    references = add_references(
        document,
        {
            "z": {"id": "z", "author": "Zeta, Ana", "title": "Z", "year": "2020"},
            "a": {"id": "a", "author": "Ávila, Bruno", "title": "A", "year": "2021"},
        },
    )

    assert len(document.tables) == 1
    assert document.tables[0].cell(0, 0).text == "Coluna A"
    assert references[0].startswith("ÁVILA")
    assert references[1].startswith("ZETA")


def test_force_ooxml_black_rewrites_theme_colors(tmp_path: Path) -> None:
    path = tmp_path / "cores.docx"
    document = Document()
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Texto")
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    document.save(path)

    # Insere artificialmente uma cor temática não preta para caracterizar o pós-processamento.
    temp = tmp_path / "mutado.docx"
    with ZipFile(path, "r") as zin, ZipFile(temp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                xml = data.decode("utf-8")
                xml = xml.replace(
                    '<w:color w:val="4472C4"/>',
                    '<w:color w:val="4472C4" w:themeColor="accent1"/>',
                    1,
                )
                data = xml.encode("utf-8")
            zout.writestr(item, data)
    temp.replace(path)

    force_ooxml_black_docx(path)

    with ZipFile(path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert 'w:val="4472C4"' not in xml
    assert 'w:themeColor=' not in xml
    assert '<w:color w:val="000000"/>' in xml


def test_validate_docx_detects_missing_sections_and_visible_residue(tmp_path: Path) -> None:
    path = tmp_path / "invalido.docx"
    document = Document()
    document.add_paragraph("Texto com \\cite{chave} e academic_pipeline:interno.")
    document.save(path)

    report = validate_docx(path)

    assert not report["ok"]
    assert any("seção de referências" in item for item in report["warnings"])
    assert any("resumo" in item for item in report["warnings"])
    assert any("comando \\cite" in item for item in report["warnings"])
    assert any("marcador técnico" in item for item in report["warnings"])


def test_render_docx_for_article_creates_docx_and_audit_jsons(tmp_path: Path) -> None:
    paths = _paths(tmp_path)
    paths.output_dir.mkdir(parents=True)
    paths.org.write_text(
        "#+TITLE: Artigo de Teste\n"
        "#+AUTHOR: Autora\n"
        "#+LATEX_HEADER: \\institution{Fundação Getúlio Vargas}\n"
        "#+LATEX_HEADER: \\coursename{Mestrado em Políticas Públicas}\n"
        "#+LATEX_HEADER: \\disciplinename{Decisões Baseadas em Evidências}\n"
        "#+LATEX_HEADER: \\professorname{Professor X}\n"
        "* Introdução\nTexto curto.\n",
        encoding="utf-8",
    )
    paths.bib.write_text(_bib_text(), encoding="utf-8")
    paths.document_json.write_text(
        json.dumps(
            {
                "metadata": {"titulo": "Artigo de Teste", "ano": "2026"},
                "sections": [
                    {
                        "title": "Introdução",
                        "text": "Texto acadêmico suficientemente longo para caracterizar o corpo.",
                    },
                    {
                        "title": "Desenvolvimento",
                        "text": "Análise com \\parencite{silva2020} e discussão.",
                    },
                    {
                        "title": "Conclusão",
                        "text": "Síntese dos resultados.",
                    },
                ],
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    paths.resumos_json.write_text(
        json.dumps(
            {
                "items": {
                    "pt-br": {
                        "abstract": "Resumo acadêmico suficientemente claro.",
                        "keywords": ["política pública", "evidências"],
                    }
                }
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    paths.compliance_json.write_text("{}", encoding="utf-8")
    paths.cfg.write_text(
        '[documento]\ntitulo="Artigo de Teste"\n'
        '[atividade]\nprofessor="Professor X"\n',
        encoding="utf-8",
    )

    result = render_docx_for_article(
        tmp_path,
        prefix="artigo",
        cfg=paths.cfg,
        quiet=True,
    )

    assert result == paths.docx
    assert paths.docx.is_file()
    assert paths.canonical_json.is_file()
    report_path = paths.output_dir / "artigo.docx_canonico_report.json"
    assert report_path.is_file()

    canonical = json.loads(paths.canonical_json.read_text(encoding="utf-8"))
    report = json.loads(report_path.read_text(encoding="utf-8"))
    assert canonical["schema_version"] == "docx-canonico-v14"
    assert canonical["bibliography_entries"] == 1
    assert canonical["references_rendered"] == 1
    assert report["ok"]
    assert report["font_color_policy"] == "all_word_color_tags_forced_to_000000"
