from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.shared import Cm

from diagnostics import (
    clean_aux_files,
    make_run_report,
    validate_docx_file,
    write_outputs_manifest,
)
from document_model import (
    AbstractBlock,
    AcademicDocument,
    BibliographyInfo,
    Block,
    Citation,
    DocumentMetadata,
    Section,
    TableData,
    TextSpan,
)
from render_docx import (
    _paper_abstract_sidecar_exists,
    bib_to_reference_map,
    citation_text,
    first_author_label,
    reference_text,
    render_docx,
    render_markdown_for_pandoc,
)


def _paper_document(*, tipo: str = "paper") -> AcademicDocument:
    return AcademicDocument(
        metadata=DocumentMetadata(
            tipo_documento=tipo,
            titulo="Título de caracterização",
            autor="Autora de Teste",
            instituicao="Fundação Getúlio Vargas",
            programa="Programa de Pós-Graduação",
            curso="Mestrado em Políticas Públicas",
            turma="Turma 1",
            polo="Brasília",
            disciplina="Decisões Baseadas em Evidências",
            professor="Professor de Teste",
            cidade="Brasília",
            ano="2026",
            data="15/07/2026",
        ),
        abstract=AbstractBlock(
            texto="Resumo acadêmico do documento.",
            palavras_chave=["políticas públicas", "evidências"],
        ),
        sections=[
            Section(
                title="Introdução",
                blocks=[
                    Block(
                        type="paragraph",
                        content=[
                            TextSpan(text="Texto com "),
                            Citation(keys=["silva2020"]),
                            TextSpan(text=" e conclusão."),
                        ],
                    ),
                    Block(
                        type="table",
                        table=TableData(
                            headers=["Dimensão", "Resultado"],
                            rows=[["Eficiência", "Positivo"]],
                        ),
                    ),
                ],
            ),
            Section(
                title="Referências",
                blocks=[Block(type="paragraph", text="Não deve duplicar.")],
            ),
        ],
        bibliography=BibliographyInfo(
            bib_path="referencias.bib",
            entries_used=["silva2020"],
        ),
    )


def _write_bib(path: Path) -> None:
    path.write_text(
        """
@article{silva2020,
  author = {Silva, Ana and Souza, Bruno},
  title = {Políticas públicas baseadas em evidências},
  year = {2020},
  journaltitle = {Revista de Políticas},
  volume = {10},
  number = {2},
  pages = {1--20},
  doi = {10.1000/teste}
}
""".strip()
        + "\n",
        encoding="utf-8",
    )


def _cm(value) -> float:
    return round(value.cm, 2)


def test_bibliography_map_extracts_nested_metadata(tmp_path: Path) -> None:
    bib = tmp_path / "referencias.bib"
    _write_bib(bib)

    refs = bib_to_reference_map(bib)

    assert list(refs) == ["silva2020"]
    assert refs["silva2020"]["author"] == "Silva, Ana and Souza, Bruno"
    assert refs["silva2020"]["title"] == "Políticas públicas baseadas em evidências"
    assert refs["silva2020"]["doi"] == "10.1000/teste"


def test_first_author_label_supports_comma_and_plain_names() -> None:
    assert first_author_label("Silva, Ana and Souza, Bruno") == "Silva"
    assert first_author_label("Ana Maria Silva") == "Silva"
    assert first_author_label("") == "Autor"


def test_citation_text_materializes_parenthetical_and_narrative_modes() -> None:
    refs = {"silva2020": {"author": "Silva, Ana", "year": "2020"}}

    parenthetical = citation_text(Citation(keys=["silva2020"]), refs)
    narrative = citation_text(Citation(mode="narrative", keys=["silva2020"]), refs)

    assert parenthetical == "(Silva, 2020)"
    assert narrative == "Silva (2020)"


def test_reference_text_includes_venue_volume_pages_and_doi() -> None:
    rendered = reference_text(
        "silva2020",
        {
            "author": "Silva, Ana",
            "year": "2020",
            "title": "Título",
            "journaltitle": "Revista",
            "volume": "10",
            "number": "2",
            "pages": "1--20",
            "doi": "10.1000/teste",
        },
    )

    assert "Silva, Ana." in rendered
    assert "Revista, 10(2), 1--20." in rendered
    assert rendered.endswith("https://doi.org/10.1000/teste")


def test_markdown_for_pandoc_skips_generated_reference_section() -> None:
    markdown = render_markdown_for_pandoc(
        _paper_document(),
        Path("/tmp"),
        cfg={},
    )

    assert "# Introdução" in markdown
    assert "Não deve duplicar." not in markdown
    assert markdown.count("# REFERÊNCIAS") == 1
    assert "[@silva2020]" in markdown


def test_sidecar_detection_requires_paper_profile_and_enabled_section(tmp_path: Path) -> None:
    output = tmp_path / "artigo.docx"
    sidecar = tmp_path / "artigo.resumos_paper.json"
    sidecar.write_text("{}", encoding="utf-8")
    cfg = {
        "projeto": {"preset": "paper_local_fgv"},
        "resumos_paper": {"ativo": True, "gerar_resumo_principal": True},
        "paths": {"document_prefix": "artigo"},
    }

    assert _paper_abstract_sidecar_exists(output, cfg)
    assert not _paper_abstract_sidecar_exists(output, {"projeto": {"preset": "outro"}})


def test_render_docx_paper_preserves_contract_and_margins(tmp_path: Path) -> None:
    bib = tmp_path / "referencias.bib"
    _write_bib(bib)
    output = tmp_path / "artigo.docx"

    result = render_docx(_paper_document(), output, bib_path=bib)

    assert result == output
    assert output.is_file()
    document = Document(output)
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Título de caracterização" in text
    assert "INTRODUÇÃO" in text
    assert text.count("REFERÊNCIAS") == 1
    assert "Políticas públicas baseadas em evidências" in text
    assert len(document.tables) == 1

    section = document.sections[0]
    assert _cm(section.top_margin) == 3.0
    assert _cm(section.left_margin) == 3.0
    assert _cm(section.right_margin) == 2.0
    assert _cm(section.bottom_margin) == 2.0


def test_render_docx_activity_places_expected_title_in_technical_sheet(tmp_path: Path) -> None:
    bib = tmp_path / "referencias.bib"
    _write_bib(bib)
    output = tmp_path / "atividade.docx"

    render_docx(_paper_document(tipo="atividade"), output, bib_path=bib)

    document = Document(output)
    assert document.tables
    table_text = "\n".join(cell.text for row in document.tables[0].rows for cell in row.cells)
    assert "FICHA TÉCNICA" in "\n".join(p.text for p in document.paragraphs)
    assert "Título do trabalho" in table_text
    assert "Título de caracterização" in table_text

    validation = validate_docx_file(
        output,
        expected_title="Título de caracterização",
        require_references=True,
    )
    assert validation["ok"]


def test_render_docx_suppresses_model_abstract_when_sidecar_exists(tmp_path: Path) -> None:
    output = tmp_path / "artigo.docx"
    (tmp_path / "artigo.resumos_paper.json").write_text("{}", encoding="utf-8")
    cfg = {
        "projeto": {"preset": "paper_local_fgv"},
        "resumos_paper": {"ativo": True, "gerar_resumo_principal": True},
        "paths": {"document_prefix": "artigo"},
    }

    render_docx(_paper_document(), output, cfg=cfg)

    text = "\n".join(p.text for p in Document(output).paragraphs)
    assert "Resumo acadêmico do documento." not in text


def test_validate_docx_file_reports_missing_file(tmp_path: Path) -> None:
    result = validate_docx_file(tmp_path / "ausente.docx")

    assert not result["ok"]
    assert result["warnings"] == ["DOCX não encontrado."]


def test_validate_docx_file_finds_title_inside_table(tmp_path: Path) -> None:
    output = tmp_path / "table.docx"
    document = Document()
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Título"
    table.cell(0, 1).text = "Política Pública"
    for text in ["Parágrafo um", "Parágrafo dois", "Referências"]:
        document.add_paragraph(text)
    document.save(output)

    result = validate_docx_file(
        output,
        expected_title="Política Pública",
        require_references=True,
    )

    assert result["ok"]
    assert not result["warnings"]


def test_validate_docx_file_warns_when_references_are_required(tmp_path: Path) -> None:
    output = tmp_path / "sem_refs.docx"
    document = Document()
    for text in ["Título", "Parágrafo um", "Parágrafo dois"]:
        document.add_paragraph(text)
    document.save(output)

    result = validate_docx_file(output, expected_title="Título", require_references=True)

    assert not result["ok"]
    assert "Seção de referências não identificada no DOCX." in result["warnings"]


def test_make_run_report_normalizes_paths_and_hashes_inputs(tmp_path: Path) -> None:
    input_zip = tmp_path / "docs.zip"
    input_zip.write_bytes(b"conteudo")
    manifest = tmp_path / "doi.csv"
    manifest.write_text("arquivo,doi\nx,10.1/x\n", encoding="utf-8")
    cfg = {
        "__config_dir__": str(tmp_path),
        "__config_path__": str(tmp_path / "cfg.toml"),
        "documentos_locais": {
            "input_zip": input_zip.name,
            "doi_manifest_path": manifest.name,
        },
        "documento": {"tipo_documento": "paper"},
        "bibliografia": {"latex_style": "abnt"},
        "latex": {"pdf_engine": "lualatex"},
    }

    report = make_run_report(
        cfg=cfg,
        config_path=None,
        out_dir=tmp_path / "output",
        prefix="artigo",
        model="modelo",
        outputs={"docx": tmp_path / "output" / "artigo.docx", "pdf": None},
        warnings=["aviso"],
        extra={"mode": "test"},
    )

    assert report["input_zip_sha256"]
    assert report["doi_manifest_sha256"]
    assert report["outputs"]["docx"].endswith("artigo.docx")
    assert report["outputs"]["pdf"] is None
    assert report["warnings"] == ["aviso"]
    assert report["extra"] == {"mode": "test"}


def test_outputs_manifest_serializes_nested_sections(tmp_path: Path) -> None:
    path = tmp_path / "outputs.txt"

    write_outputs_manifest(
        path,
        {
            "org": "/tmp/a.org",
            "idiomas": {"en": "/tmp/en.docx", "es": None},
        },
    )

    assert path.read_text(encoding="utf-8") == (
        "org: /tmp/a.org\n"
        "[idiomas]\n"
        "en: /tmp/en.docx\n"
        "es: None\n"
    )


def test_clean_aux_files_removes_only_declared_artifacts(tmp_path: Path) -> None:
    org = tmp_path / "artigo.org"
    org.write_text("* Texto\n", encoding="utf-8")
    removable = [tmp_path / "artigo.aux", tmp_path / "artigo.log", tmp_path / "outro.tex"]
    for path in removable:
        path.write_text("x", encoding="utf-8")
    preserved = tmp_path / "artigo.bib"
    preserved.write_text("x", encoding="utf-8")

    removed = clean_aux_files(org)

    assert set(removed) == {str(path) for path in removable}
    assert preserved.exists()
    assert org.exists()


def test_rendered_docx_is_valid_ooxml_package(tmp_path: Path) -> None:
    output = tmp_path / "artigo.docx"
    render_docx(_paper_document(), output)

    with ZipFile(output) as archive:
        names = set(archive.namelist())

    assert "[Content_Types].xml" in names
    assert "word/document.xml" in names
    assert "word/styles.xml" in names
