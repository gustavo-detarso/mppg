#!/usr/bin/env python3
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt

out = Path(__file__).with_name('reference_fgv.docx')
doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(3)
sec.left_margin = Cm(3)
sec.right_margin = Cm(2)
sec.bottom_margin = Cm(2)
styles = doc.styles
for name in ['Normal', 'Title', 'Heading 1', 'Heading 2', 'Heading 3']:
    st = styles[name]
    st.font.name = 'Times New Roman'
    st.font.size = Pt(12 if name == 'Normal' else 14)
    if name.startswith('Heading') or name == 'Title':
        st.font.bold = True
p = doc.add_paragraph('Fundação Getúlio Vargas')
p.alignment = 1
p = doc.add_paragraph('Modelo de referência FGV para DOCX gerado pelo academic_pipeline rc10')
p.alignment = 1
doc.add_page_break()
doc.add_heading('Título de seção', level=1)
p = doc.add_paragraph('Parágrafo de exemplo com espaçamento e margens acadêmicas.')
p.paragraph_format.line_spacing = 1.5
p.paragraph_format.first_line_indent = Cm(1.25)
doc.save(out)
print(out)
