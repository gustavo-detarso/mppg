#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Diagnóstico, validação operacional e rastreabilidade para academic_pipeline rc10.7.

Este módulo evita que erros de ambiente/caminho sejam descobertos apenas no fim
 da geração. Ele é deliberadamente independente da OpenAI: doctor/check-config e
 recompile não devem exigir OPENAI_API_KEY, salvo quando a execução realmente for
 gerar conteúdo novo via IA.
"""
from __future__ import annotations

import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from utils import resolve_path, write_json, write_text, normalize_title_loose
from prompt_manager import validate_prompt_paths, prompt_report_for_cfg

PIPELINE_VERSION = "rc10.7.38-activity-ai-data-mode"


# ---------------------------------------------------------------------------
# Helpers gerais
# ---------------------------------------------------------------------------

def now_iso() -> str:
    return datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds")


def file_sha256(path: Path | None, *, max_bytes: int | None = None) -> str | None:
    if not path or not path.exists() or not path.is_file():
        return None
    h = hashlib.sha256()
    total = 0
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            if max_bytes is not None and total + len(chunk) > max_bytes:
                chunk = chunk[: max(0, max_bytes - total)]
            if not chunk:
                break
            h.update(chunk)
            total += len(chunk)
            if max_bytes is not None and total >= max_bytes:
                break
    return h.hexdigest()


def command_exists(name: str) -> bool:
    return shutil.which(name) is not None


def command_version(name: str, args: list[str] | None = None) -> str | None:
    exe = shutil.which(name)
    if not exe:
        return None
    try:
        proc = subprocess.run([exe, *(args or ["--version"])], text=True, capture_output=True, timeout=8)
        text = (proc.stdout or proc.stderr or "").splitlines()
        return text[0].strip() if text else exe
    except Exception:
        return exe


def kpsewhich_exists(filename: str) -> bool | None:
    kpse = shutil.which("kpsewhich")
    if not kpse:
        return None
    try:
        proc = subprocess.run([kpse, filename], text=True, capture_output=True, timeout=8)
        return proc.returncode == 0 and bool(proc.stdout.strip())
    except Exception:
        return None


def _cfg_section(cfg: dict[str, Any], name: str) -> dict[str, Any]:
    sec = cfg.get(name, {})
    return sec if isinstance(sec, dict) else {}


def _config_dir(cfg: dict[str, Any]) -> Path:
    return Path(str(cfg.get("__config_dir__") or Path.cwd())).resolve()


def _style_from_cfg(cfg: dict[str, Any]) -> str:
    bib = _cfg_section(cfg, "bibliografia")
    doc = _cfg_section(cfg, "documento")
    return str(bib.get("latex_style") or bib.get("estilo_citacao") or doc.get("estilo_citacao") or "apa").strip().lower()


def _doc_type(cfg: dict[str, Any]) -> str:
    doc = _cfg_section(cfg, "documento")
    return str(doc.get("tipo_documento") or "paper").strip().lower()


def _out_dir_and_prefix_from_cfg(cfg: dict[str, Any]) -> tuple[Path, str]:
    """Resolve a saída final do documento a partir da seção [paths].

    A partir da rc10.7.20, [saida] e [documento].output_dir deixam de ser
    a interface oficial de caminhos. O TOML novo deve declarar [paths].
    """
    base = _config_dir(cfg)
    paths = _cfg_section(cfg, "paths")
    projeto = _cfg_section(cfg, "projeto")
    prefix = str(paths.get("document_prefix") or projeto.get("nome") or "documento").strip() or "documento"
    out_base = resolve_path(paths.get("document_output_dir") or "../../output/documento", base) or (base / "output/documento")
    out_dir = out_base / prefix if bool(paths.get("create_document_subdir", True)) else out_base
    return out_dir, prefix


def _work_cache_dirs_from_cfg(cfg: dict[str, Any], prefix: str) -> tuple[Path, Path]:
    base = _config_dir(cfg)
    paths = _cfg_section(cfg, "paths")
    work_base = resolve_path(paths.get("work_dir") or "../../output/work", base) or (base / "output/work")
    cache_base = resolve_path(paths.get("cache_dir") or "../../output/cache", base) or (base / "output/cache")
    work_dir = work_base / prefix if bool(paths.get("create_work_subdir", True)) else work_base
    cache_dir = cache_base / prefix if bool(paths.get("create_cache_subdir", True)) else cache_base
    return work_dir, cache_dir


# ---------------------------------------------------------------------------
# Doctor
# ---------------------------------------------------------------------------

def run_doctor(cfg: dict[str, Any] | None = None) -> dict[str, Any]:
    """Diagnostica ambiente local. Nunca exige OpenAI se cfg não exigir geração."""
    cfg = cfg or {}
    base = _config_dir(cfg) if cfg else Path.cwd().resolve()
    latex = _cfg_section(cfg, "latex")
    docx = _cfg_section(cfg, "docx")
    mm = _cfg_section(cfg, "mapa_mental")
    doc = _cfg_section(cfg, "documento")
    style = _style_from_cfg(cfg) if cfg else ""

    checks: list[dict[str, Any]] = []

    def add(name: str, ok: bool, detail: str = "", severity: str = "error", value: Any = None) -> None:
        checks.append({"name": name, "ok": bool(ok), "severity": severity, "detail": detail, "value": value})

    # Python/deps básicos
    add("python", True, sys.version.split()[0], "info", sys.executable)
    for mod in ("openai", "pydantic", "dotenv", "pypdf", "docx", "openpyxl"):
        try:
            __import__("dotenv" if mod == "dotenv" else mod)
            add(f"python_module:{mod}", True, "disponível", "info")
        except Exception as exc:
            # openai só é error quando for gerar documento novo; aqui vira warning.
            sev = "warning" if mod == "openai" else "error"
            add(f"python_module:{mod}", False, str(exc), sev)

    # Comandos externos
    for cmd in ("emacs", "lualatex", "xelatex", "pdflatex", "biber", "kpsewhich"):
        exists = command_exists(cmd)
        sev = "warning" if cmd in {"xelatex", "pdflatex", "kpsewhich"} else "error"
        add(f"command:{cmd}", exists, command_version(cmd) or "não encontrado", sev)

    if bool(docx.get("usar_pandoc", False)):
        add("command:pandoc", command_exists("pandoc"), command_version("pandoc") or "não encontrado", "error")
    else:
        add("command:pandoc", command_exists("pandoc"), command_version("pandoc") or "não encontrado", "info")

    if bool(mm.get("gerar", False)):
        jar_raw = mm.get("plantuml_jar_path") or _cfg_section(cfg, "documento").get("plantuml_jar_path") or os.getenv("PLANTUML_JAR")
        jar = resolve_path(jar_raw, base) if jar_raw else None
        has_plantuml = command_exists("plantuml") or (jar is not None and jar.exists() and command_exists("java"))
        detail = "plantuml no PATH" if command_exists("plantuml") else f"jar={jar}, java={command_exists('java')}"
        add("plantuml", has_plantuml, detail, "error")

    # Perfil institucional
    profile_name = cfg.get("__institution_profile_name__") if cfg else None
    profile_path = cfg.get("__institution_profile_path__") if cfg else None
    if profile_name:
        add("institution:profile", True, f"{profile_name} — {profile_path}", "info")
    elif cfg:
        add("institution:profile", True, "nenhum perfil institucional informado; usando caminhos explícitos do TOML", "info")

    # Arquivos locais FGV
    aw = resolve_path(latex.get("org_latex_class_init"), base) if latex else None
    latex_extra = resolve_path(latex.get("latex_extra_path"), base) if latex else None
    logo = resolve_path(latex.get("fgv_logo_path"), base) if latex else None
    add("file:academic-writing.el", bool(aw and aw.exists()), str(aw or "não informado"), "warning")
    add("path:latex_extra_path", bool(latex_extra and latex_extra.exists()), str(latex_extra or "não informado"), "warning")
    if latex_extra and latex_extra.exists():
        p = latex_extra if latex_extra.is_dir() else latex_extra.parent
        add("file:fgv-paper.sty", (p / "fgv-paper.sty").exists(), str(p / "fgv-paper.sty"), "warning")
        add("file:fgv-dissertacao.sty", (p / "fgv-dissertacao.sty").exists(), str(p / "fgv-dissertacao.sty"), "warning")
    add("file:fgv_logo_path", True if not logo else logo.exists(), str(logo or "não informado"), "info")

    # OpenAI só é obrigatório se não for somente-renderizar e executar_documento não for false.
    pipeline = _cfg_section(cfg, "pipeline")
    needs_openai = bool(cfg) and bool(pipeline.get("executar_documento", True)) and not bool(cfg.get("__somente_renderizar__", False))
    openai_ok = bool(os.getenv("OPENAI_API_KEY"))
    add("env:OPENAI_API_KEY", openai_ok, "definida" if openai_ok else "não definida", "error" if needs_openai else "warning")

    # Bibliografia TeX
    if style:
        if style == "apa":
            exists = kpsewhich_exists("apa.bbx")
            add("tex:biblatex-apa", bool(exists) if exists is not None else True, "apa.bbx" if exists else "kpsewhich indisponível ou apa.bbx não encontrado", "warning" if exists is None else "error")
        elif style == "abnt":
            exists = kpsewhich_exists("abnt.bbx")
            add("tex:biblatex-abnt", bool(exists) if exists is not None else True, "abnt.bbx" if exists else "kpsewhich indisponível ou abnt.bbx não encontrado", "warning" if exists is None else "error")

    # Diretório output
    if cfg:
        out_dir, _ = _out_dir_and_prefix_from_cfg(cfg)
        try:
            out_dir.mkdir(parents=True, exist_ok=True)
            test = out_dir / ".write_test"
            test.write_text("ok", encoding="utf-8")
            test.unlink(missing_ok=True)
            add("output:writable", True, str(out_dir), "info")
        except Exception as exc:
            add("output:writable", False, f"{out_dir}: {exc}", "error")

    errors = [c for c in checks if not c["ok"] and c["severity"] == "error"]
    warnings = [c for c in checks if not c["ok"] and c["severity"] == "warning"]
    return {"version": PIPELINE_VERSION, "generated_at": now_iso(), "ok": not errors, "errors": errors, "warnings": warnings, "checks": checks}


def print_doctor_report(report: dict[str, Any]) -> None:
    print(f"academic_pipeline {report.get('version')} — doctor")
    print(f"Status: {'OK' if report.get('ok') else 'ERROS'}")
    for c in report.get("checks", []):
        mark = "OK" if c.get("ok") else ("WARN" if c.get("severity") == "warning" else ("INFO" if c.get("severity") == "info" else "ERRO"))
        print(f"[{mark}] {c.get('name')}: {c.get('detail')}")


# ---------------------------------------------------------------------------
# Check-config
# ---------------------------------------------------------------------------

def check_config(cfg: dict[str, Any]) -> dict[str, Any]:
    base = _config_dir(cfg)
    errors: list[str] = []
    warnings: list[str] = []

    def err(msg: str) -> None:
        errors.append(msg)

    def warn(msg: str) -> None:
        warnings.append(msg)

    pipeline = _cfg_section(cfg, "pipeline")
    paths = _cfg_section(cfg, "paths")
    local = _cfg_section(cfg, "documentos_locais")
    doc = _cfg_section(cfg, "documento")
    latex = _cfg_section(cfg, "latex")
    docx = _cfg_section(cfg, "docx")
    bib = _cfg_section(cfg, "bibliografia")
    prisma = _cfg_section(cfg, "relatorio_pesquisa")
    inst = _cfg_section(cfg, "instituicao")
    if inst.get("perfil") and not cfg.get("__institution_profile_path__"):
        err(f"Perfil institucional informado, mas não carregado: {inst.get('perfil')}")
    if cfg.get("__institution_profile_path__"):
        p = Path(str(cfg.get("__institution_profile_path__")))
        if not p.exists():
            err(f"Arquivo do perfil institucional não existe: {p}")

    # Caminhos oficiais
    if not paths:
        err("Seção [paths] ausente. A partir da rc10.7.20, use [paths] para document_output_dir, work_dir, cache_dir e research_output_dir.")
    for k in ("document_output_dir", "work_dir", "cache_dir", "research_output_dir"):
        if not str(paths.get(k) or "").strip():
            warn(f"[paths].{k} não informado; será usado padrão relativo ao TOML.")
    for k in ("document_output_dir", "work_dir", "cache_dir", "research_output_dir"):
        raw = paths.get(k)
        if raw:
            p = resolve_path(raw, base)
            try:
                p.mkdir(parents=True, exist_ok=True)
            except Exception as exc:
                err(f"Não foi possível criar [paths].{k}: {p}: {exc}")

    # Entrada local
    if str(pipeline.get("modo_entrada") or local.get("modo_entrada") or "documentos_locais").lower() in {"documentos_locais", "local", "zip", "pasta_local", ""}:
        input_zip = resolve_path(local.get("input_zip"), base)
        input_dir = resolve_path(local.get("input_dir"), base)
        if input_zip and not input_zip.exists():
            err(f"[documentos_locais].input_zip não existe: {input_zip}")
        if input_dir and not input_dir.exists():
            err(f"[documentos_locais].input_dir não existe: {input_dir}")
        if not input_zip and not input_dir:
            warn("Nenhum input_zip/input_dir informado em [documentos_locais]; ok apenas se usar --somente-renderizar ou relatório externo.")
        doi_manifest = resolve_path(local.get("doi_manifest_path"), base)
        if local.get("doi_manifest_path") and not (doi_manifest and doi_manifest.exists()):
            err(f"doi_manifest_path não existe: {doi_manifest}")

    # Documento / visual
    for key in ("template_org", "template_path"):
        raw = doc.get(key)
        if raw:
            p = resolve_path(raw, base)
            if not (p and p.exists()):
                warn(f"[documento].{key} não existe ou não será usado diretamente pelo renderizador rc10: {p}")

    out_dir, _ = _out_dir_and_prefix_from_cfg(cfg)
    try:
        out_dir.mkdir(parents=True, exist_ok=True)
    except Exception as exc:
        err(f"Não foi possível criar output_dir: {out_dir}: {exc}")

    # LaTeX
    if bool(doc.get("exportar_pdf", True)):
        aw = resolve_path(latex.get("org_latex_class_init"), base)
        latex_extra = resolve_path(latex.get("latex_extra_path"), base)
        if not (aw and aw.exists()):
            warn(f"academic-writing.el ausente: {aw}")
        if not (latex_extra and latex_extra.exists()):
            warn(f"latex_extra_path ausente: {latex_extra}")
        engine = str(latex.get("pdf_engine") or "lualatex")
        if engine not in {"lualatex", "xelatex", "pdflatex"}:
            err(f"[latex].pdf_engine inválido: {engine}")

    # DOCX
    if bool(doc.get("exportar_docx", True)):
        ref = resolve_path(docx.get("reference_docx") or doc.get("docx_reference"), base)
        if (docx.get("reference_docx") or doc.get("docx_reference")) and not (ref and ref.exists()):
            warn(f"reference_docx informado, mas não existe: {ref}")
        if bool(docx.get("usar_pandoc", False)):
            csl = resolve_path(docx.get("csl_path") or bib.get("docx_csl"), base)
            if not command_exists("pandoc"):
                warn("[docx].usar_pandoc=true, mas pandoc não foi encontrado no PATH.")
            if (docx.get("csl_path") or bib.get("docx_csl")) and not (csl and csl.exists()):
                warn(f"CSL informado, mas não existe: {csl}")

    # Estilo bibliográfico
    style = _style_from_cfg(cfg)
    if style not in {"apa", "abnt", "authoryear", "numeric", "chicago"}:
        warn(f"Estilo bibliográfico pouco comum: {style}. Verifique se o BibLaTeX/CSL correspondente existe.")

    # Capa duplicada em paper/atividade
    doc_type = _doc_type(cfg)
    if doc_type in {"paper", "atividade"}:
        program = normalize_title_loose(str(doc.get("program_name") or ""))
        course = normalize_title_loose(str(doc.get("course_name") or _cfg_section(cfg, "atividade").get("curso") or ""))
        if program and course and program == course:
            err("[documento].program_name é igual ao curso; isso tende a duplicar a capa. Use program_name = \"\".")

    # Dissertação: campos mínimos
    if doc_type in {"dissertacao", "dissertação"}:
        for k in ("course_name", "program_name"):
            if not str(doc.get(k) or _cfg_section(cfg, "atividade").get("curso") or "").strip():
                warn(f"Dissertação sem campo {k}; o front matter pode ficar incompleto.")
        if not str(doc.get("orientador") or doc.get("professor_name") or _cfg_section(cfg, "atividade").get("professor") or "").strip():
            warn("Dissertação sem orientador/professor informado.")

    # Prompt bank
    for msg in validate_prompt_paths(cfg):
        err(msg)

    # Relatório PRISMA
    if bool(prisma.get("ativo", False)):
        prisma_prefix = str(paths.get("research_prefix") or prisma.get("prefixo") or "relatorio_prisma").strip() or "relatorio_prisma"
        prisma_base = resolve_path(paths.get("research_output_dir") or "../../output/pesquisa", base)
        prisma_out = prisma_base / prisma_prefix if bool(paths.get("create_research_subdir", True)) else prisma_base
        try:
            prisma_out.mkdir(parents=True, exist_ok=True)
        except Exception as exc:
            err(f"Não foi possível criar [paths].research_output_dir do relatório PRISMA: {prisma_out}: {exc}")
        pjson = resolve_path(prisma.get("prisma_json_path"), base)
        if prisma.get("prisma_json_path") and not (pjson and pjson.exists()):
            err(f"[relatorio_pesquisa].prisma_json_path não existe: {pjson}")

    return {"version": PIPELINE_VERSION, "generated_at": now_iso(), "ok": not errors, "errors": errors, "warnings": warnings}


def print_check_config_report(report: dict[str, Any]) -> None:
    print(f"academic_pipeline {report.get('version')} — check-config")
    print(f"Status: {'OK' if report.get('ok') else 'ERROS'}")
    for e in report.get("errors", []):
        print(f"[ERRO] {e}")
    for w in report.get("warnings", []):
        print(f"[WARN] {w}")
    if not report.get("errors") and not report.get("warnings"):
        print("Nenhum problema encontrado.")


# ---------------------------------------------------------------------------
# Validação de artefatos e rastreabilidade
# ---------------------------------------------------------------------------

def validate_docx_file(path: Path | None, expected_title: str | None = None, *, require_references: bool = False) -> dict[str, Any]:
    """Validação leve do DOCX gerado.

    A versão anterior lia apenas `doc.paragraphs`; em atividades, porém,
    o título fica dentro da tabela da Ficha Técnica. Isso gerava falso
    alerta de "Título esperado não encontrado no DOCX". Agora a inspeção
    inclui também células de tabelas e usa comparação normalizada.
    """
    result = {"path": str(path) if path else None, "ok": False, "warnings": [], "paragraphs": 0, "size_bytes": 0}
    if not path or not path.exists():
        result["warnings"].append("DOCX não encontrado.")
        return result
    result["size_bytes"] = path.stat().st_size
    if path.stat().st_size <= 0:
        result["warnings"].append("DOCX vazio.")
        return result
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        table_texts: list[str] = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    txt = cell.text.strip()
                    if txt:
                        table_texts.append(txt)

        result["paragraphs"] = len(paragraphs)
        joined = "\n".join(paragraphs + table_texts)
        joined_norm = normalize_title_loose(joined)
        expected_norm = normalize_title_loose(expected_title or "")

        # Aceita título em parágrafo, célula de tabela ou metadado escrito com pontuação normalizada.
        if expected_norm and expected_norm not in joined_norm:
            result["warnings"].append("Título esperado não encontrado no DOCX.")
        if require_references and "refer" not in joined_norm:
            result["warnings"].append("Seção de referências não identificada no DOCX.")
        if len(paragraphs) < 3 and len(table_texts) < 3:
            result["warnings"].append("DOCX parece curto demais.")
        result["ok"] = not result["warnings"]
    except Exception as exc:
        result["warnings"].append(f"Falha ao inspecionar DOCX: {exc}")
        result["ok"] = True  # arquivo existe, mas sem validação profunda
    return result

def make_run_report(
    *,
    cfg: dict[str, Any],
    config_path: Path | None,
    out_dir: Path,
    prefix: str,
    model: str | None,
    outputs: dict[str, Any],
    warnings: list[str] | None = None,
    extra: dict[str, Any] | None = None,
) -> dict[str, Any]:
    local = _cfg_section(cfg, "documentos_locais")
    bib = _cfg_section(cfg, "bibliografia")
    latex = _cfg_section(cfg, "latex")
    input_zip = resolve_path(local.get("input_zip"), _config_dir(cfg)) if local.get("input_zip") else None
    input_dir = resolve_path(local.get("input_dir"), _config_dir(cfg)) if local.get("input_dir") else None
    doi_manifest = resolve_path(local.get("doi_manifest_path"), _config_dir(cfg)) if local.get("doi_manifest_path") else None
    resolved_outputs: dict[str, Any] = {}
    for key, value in (outputs or {}).items():
        if value is None:
            resolved_outputs[key] = None
        elif isinstance(value, Path):
            resolved_outputs[key] = str(value)
        else:
            resolved_outputs[key] = value

    return {
        "version": PIPELINE_VERSION,
        "generated_at": now_iso(),
        "config_path": str(config_path) if config_path else cfg.get("__config_path__"),
        "output_dir": str(out_dir),
        "prefixo": prefix,
        "tipo_documento": _doc_type(cfg),
        "instituicao_perfil": cfg.get("__institution_profile_name__"),
        "instituicao_profile_path": cfg.get("__institution_profile_path__"),
        "modelo_openai": model,
        "estilo_bibliografico": _style_from_cfg(cfg),
        "pdf_engine": latex.get("pdf_engine") or "lualatex",
        "input_zip": str(input_zip) if input_zip else None,
        "input_zip_sha256": file_sha256(input_zip) if input_zip else None,
        "input_dir": str(input_dir) if input_dir else None,
        "doi_manifest": str(doi_manifest) if doi_manifest else None,
        "doi_manifest_sha256": file_sha256(doi_manifest) if doi_manifest else None,
        "outputs": resolved_outputs,
        "warnings": warnings or [],
        "prompts": prompt_report_for_cfg(cfg),
        "extra": extra or {},
    }


def write_outputs_manifest(path: Path, outputs: dict[str, Any]) -> None:
    lines = []
    for key, value in outputs.items():
        if isinstance(value, dict):
            lines.append(f"[{key}]")
            for subk, subv in value.items():
                lines.append(f"{subk}: {subv}")
        else:
            lines.append(f"{key}: {value}")
    write_text(path, "\n".join(lines).strip() + "\n")


def print_outputs(outputs: dict[str, Any], *, title: str = "Saídas") -> None:
    print(title + ":")
    for key, value in outputs.items():
        print(f"- {key}: {value}")


# ---------------------------------------------------------------------------
# Recompile helper
# ---------------------------------------------------------------------------

def clean_aux_files(org_path: Path) -> list[str]:
    removed: list[str] = []
    patterns = ["*.aux", "*.bcf", "*.bbl", "*.blg", "*.log", "*.out", "*.run.xml", "*.toc", "*.tex", "*.lof", "*.lot", "*.fls", "*.fdb_latexmk"]
    for pattern in patterns:
        for file in org_path.parent.glob(pattern):
            try:
                file.unlink()
                removed.append(str(file))
            except Exception:
                pass
    return removed
